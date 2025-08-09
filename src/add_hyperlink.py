import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from docx import Document
from docx.oxml import OxmlElement

def justify_and_set_line_spacing(doc):
    try:
        for i, paragraph in enumerate(doc.paragraphs):
            if i != 1:  # Đảm bảo rằng paragraph số 2 (index 1) không bị thay đổi
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = Pt(18)
        
        return doc
    
    except Exception as e:
        return f"Lỗi xử lý căn đều và giãn cách dòng: {e}"

def check_images_in_document(doc):
    """Kiểm tra sự tồn tại của tên ảnh và in ra dòng số của đoạn chứa tên ảnh."""
    image_positions = []  # Lưu các vị trí của các ảnh trong tài liệu
    image_names = []      # Lưu tên ảnh (tên đoạn văn chứa tên ảnh)
    
    for idx, para in enumerate(doc.paragraphs):
        # Kiểm tra nếu đoạn văn là căn giữa
        if para.alignment == 1 or 1:  # 1 là mã của alignment=center
            # Kiểm tra nếu đoạn văn có chứa run in nghiêng và đoạn văn không rỗng
            for run in para.runs:
                if run.font.italic and para.text.strip():  # Kiểm tra in nghiêng và đoạn văn không rỗng
                    image_positions.append(idx + 1)  # Lưu dòng số chứa ảnh (dòng bắt đầu từ 1)
                    image_names.append(para.text.strip())  # Lưu tên ảnh (văn bản của đoạn)
                    break  # Thoát khỏi vòng lặp sau khi tìm thấy lần đầu tiên

    return image_positions, image_names


def format_image_names_in_document(doc):
    """Tìm và định dạng tên ảnh trong tài liệu."""
    # Lấy các vị trí và tên ảnh từ hàm check_images_in_document
    image_positions, image_names = check_images_in_document(doc)
    
    for idx, para in zip(image_positions, image_names):
        # Tìm đoạn văn theo vị trí
        paragraph = doc.paragraphs[idx - 1]  # Vị trí trong doc.paragraphs bắt đầu từ 0, nên trừ 1
        
        # Chuyển văn bản thành viết thường, chỉ viết hoa chữ cái đầu
        formatted_text = para.capitalize()  # Sử dụng capitalize để viết hoa chữ đầu
        
        # Cập nhật lại văn bản của đoạn
        paragraph.text = formatted_text
        
        # Định dạng đoạn văn:
        paragraph.alignment = 1  # Căn giữa (center alignment)
        
        # Định dạng các run trong đoạn văn
        for run in paragraph.runs:
            run.font.italic = True  # In nghiêng
            run.font.size = Pt(11)   # Cỡ chữ 11
    
    return doc


def clear_paragraph(paragraph):
    """
    Xóa tất cả các runs trong đoạn văn.
    """
    p = paragraph._element
    for child in p[:]:
        p.remove(child)



def format_keyword_in_paragraph(paragraph, keyword, bold=True, color=RGBColor(0, 0, 255)):
    """
    Định dạng tất cả các xuất hiện của từ khóa trong đoạn văn với font Arial, cỡ 11, giãn dòng 1.5, căn đều.
    Từ khóa sẽ có màu xanh và in đậm nếu yêu cầu.
    """
    # Đặt định dạng chung cho đoạn văn
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(11)
    paragraph.paragraph_format.line_spacing = Pt(15)  # Giãn dòng 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Căn đều

    # Lấy toàn bộ văn bản của đoạn văn
    full_text = paragraph.text
    # Tìm tất cả các vị trí xuất hiện của từ khóa (không phân biệt chữ hoa chữ thường)
    matches = list(re.finditer(re.escape(keyword), full_text, re.IGNORECASE))

    if not matches:
        return

    # Xóa tất cả các runs hiện tại
    paragraph.clear()

    last_index = 0
    for match in matches:
        start, end = match.span()
        # Thêm văn bản trước từ khóa
        if start > last_index:
            before_text = full_text[last_index:start]
            run = paragraph.add_run(before_text)
            # Đặt định dạng chung cho run này
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        
        # Thêm từ khóa với định dạng đặc biệt
        keyword_text = full_text[start:end]
        keyword_run = paragraph.add_run(keyword_text)
        keyword_run.bold = bold
        keyword_run.font.color.rgb = color
        keyword_run.font.name = 'Arial'
        keyword_run.font.size = Pt(11)
        
        last_index = end

    # Thêm văn bản còn lại sau từ khóa cuối cùng
    if last_index < len(full_text):
        after_text = full_text[last_index:]
        run = paragraph.add_run(after_text)
        # Đặt định dạng chung cho run này
        run.font.name = 'Arial'
        run.font.size = Pt(11)


def change_font_to_arial(doc):
    """
    Changes the font of all text in a docx document to Arial, preserving other formatting.

    Args:
        doc: A docx.Document object (the opened docx file).

    Returns:
        A docx.Document object with the font changed to Arial.  Returns None if an error occurs.
    """
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Preserve font size and other formatting
                original_size = run.font.size
                original_bold = run.font.bold
                original_italic = run.font.italic
                original_underline = run.font.underline
                original_color = run.font.color.rgb

                run.font.name = 'Arial'  #Set font to Arial

                # Restore other formatting attributes
                if original_size:
                    run.font.size = original_size
                if original_bold:
                    run.font.bold = True
                if original_italic:
                    run.font.italic = True
                if original_underline:
                    run.font.underline = original_underline
                if original_color:
                    run.font.color.rgb = original_color


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Preserve font size and other formatting (same as above)
                            original_size = run.font.size
                            original_bold = run.font.bold
                            original_italic = run.font.italic
                            original_underline = run.font.underline
                            original_color = run.font.color.rgb

                            run.font.name = 'Arial' #Set font to Arial

                            # Restore other formatting attributes
                            if original_size:
                                run.font.size = original_size
                            if original_bold:
                                run.font.bold = True
                            if original_italic:
                                run.font.italic = True
                            if original_underline:
                                run.font.underline = original_underline
                            if original_color:
                                run.font.color.rgb = original_color

        return doc
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """
    Helper function to add a hyperlink to a paragraph with formatting.

    Parameters:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to modify.
        url (str): The hyperlink URL.
        text (str): The text to display as the hyperlink.

    Returns:
        None
    """
    # Create a new relationship ID for the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the hyperlink text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style the text: bold, underline, blue color
    bold = OxmlElement('w:b')
    rPr.append(bold)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Set the font to Arial
    font = OxmlElement('w:sz')
    font.set(qn('w:val'), '22')  # Set font size (12pt for Arial is size 24, adjust if needed)
    rPr.append(font)

    # Add the font family as Arial
    font_family = OxmlElement('w:rFonts')
    font_family.set(qn('w:ascii'), 'Arial')
    font_family.set(qn('w:eastAsia'), 'Arial')
    font_family.set(qn('w:hAnsi'), 'Arial')
    rPr.append(font_family)

    # Style the text as a hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    run.append(rPr)

    # Add the text element
    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)

    # Append the run to the hyperlink
    hyperlink.append(run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)


import unicodedata

def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return ''.join([c for c in nfkd_form if not unicodedata.combining(c)])


def add_hyperlink_to_keyword_in_last_paragraph(doc, keyword, link, num_paragraph):
    """
    Finds a keyword (even with spaces) in a paragraph of a Word document and adds a hyperlink to it with formatting.

    Parameters:
        doc (docx.document.Document): The Word document object.
        keyword (str): The keyword to search for.
        link (str): The URL to link to.
        num_paragraph (int): Index of the paragraph to modify.

    Returns:
        docx.document.Document: Modified Word document object.
    """
    # Get the paragraph at num_paragraph index
    paragraph = doc.paragraphs[num_paragraph]

    # Normalize the text to remove accents and handle spaces
    text_lower = remove_accents(paragraph.text.lower()).strip()
    keyword_lower = remove_accents(keyword.lower()).strip()

    # Print for debugging purposes

    # Ensure we handle multi-word keywords
    if keyword_lower in text_lower:
        # Split the text into parts based on the keyword
        start_index = text_lower.find(keyword_lower)
        end_index = start_index + len(keyword)

        # Text before and after the keyword, including the space after the keyword
        before_keyword = paragraph.text[:start_index]
        after_keyword = paragraph.text[end_index:]

        # Add a space after the keyword if it exists in the original text
        if after_keyword and after_keyword[0] != ' ':
            after_keyword = ' ' + after_keyword

        # Clear the paragraph while retaining its format
        for run in paragraph.runs:
            run.clear()

        # Add text before the keyword
        if before_keyword:
            paragraph.add_run(before_keyword)

        # Add the hyperlink for the keyword
        add_hyperlink(paragraph, link, keyword)

        # Add text after the keyword
        if after_keyword:
            paragraph.add_run(after_keyword)

    else:
        print(f"Keyword not found: {keyword_lower}")

    return doc


def add_hyper_link_folder(docx_folder, text_data):
    """Xử lý các file DOCX và chèn hyperlink vào key tương ứng."""
    # Đọc dữ liệu từ biến text_data và lưu thông tin vào dictionary
    file_info = {}
    for line in text_data.strip().split("\n"):

        parts = re.split(r'[,\t]', line.strip())  # Tách theo cả dấu phẩy và tab
        
        if len(parts) == 4:  # Đảm bảo mỗi dòng có đủ 4 phần
            file_name, key, link, stt = parts
            key = key.strip()
            link = link.strip()
            stt = stt.strip()


            base_name = os.path.splitext(file_name)[0]  # Bỏ phần mở rộng nếu có
            if base_name not in file_info:
                file_info[base_name] = []
            file_info[base_name].append((key, link, stt))

    # Duyệt qua các file docx trong folder
    for file_name in os.listdir(docx_folder):
        if file_name.endswith('.docx'):
            base_name = os.path.splitext(file_name)[0]
            if base_name in file_info:

                file_path = os.path.join(docx_folder, file_name)
                doc = Document(file_path)

                # Lấy đoạn paragraph cuối cùng
                process_paragraph = doc.paragraphs[int(file_info[base_name][0][2])]
                keyword = file_info[base_name][0][0]

                print(f'Đang xử lý: {file_name}')


                if keyword.lower() not in process_paragraph.text.lower():
                    prompt = f"""Viết lại dùm tôi đoạn sau: {process_paragraph.text}. 
                            Keyword là: {keyword}. 
                            Đảm bảo 100% có duy nhất 1 keyword trong đoạn. Không được phép thay đổi nội dung của đoạn.
                            Ráng chèn 1 cách tự nhiên chút, tối đa 400 kí tự. 
                            Thông tin key liên quan nhà cái và cổng game giải trí, có thể viết là giải trí tại, giải tỏa căng thẳng ...
                            Không markdown. Chỉ gửi về kết quả tốt nhất"""
                    
                    from components.seo_write import DEFAULT_NORMALIZE_MODEL_NAME, DEFAULT_NORMALIZE_MODEL_TYPE
                    from src.chat_bot import call_chatbot
                    from src.controller import gpt_key1

                    processed_content = call_chatbot(prompt, DEFAULT_NORMALIZE_MODEL_NAME, DEFAULT_NORMALIZE_MODEL_TYPE, gpt_key1)

                    process_paragraph.text = processed_content

                # Chèn hyperlink vào keyword trong đoạn cuối
                doc = add_hyperlink_to_keyword_in_last_paragraph(doc, keyword, file_info[base_name][0][1], int(file_info[base_name][0][2]))

                # Căn chỉnh và xử lý định dạng
                doc = justify_and_set_line_spacing(doc)
                doc = format_image_names_in_document(doc)
                doc = change_font_to_arial(doc)

                # Ghi đè file cũ
                doc.save(file_path)
