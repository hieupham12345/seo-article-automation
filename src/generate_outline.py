import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
import re
from src.chat_bot import call_chatbot

def clean_json(data):
    try:
        # Chuyển thành JSON string rồi parse lại để kiểm tra
        json_str = json.dumps(data, ensure_ascii=False)
        cleaned_data = json.loads(json_str)
        return cleaned_data
    except json.JSONDecodeError as e:
        print("Lỗi JSON:", e)
        return None
    
def convert_to_dict(raw_data):
    """
    Robust function to convert a string to a dictionary, handling various markdown formats.
    
    Args:
        raw_data (str): The input string potentially containing markdown or extra formatting
    
    Returns:
        dict or None: Parsed dictionary or None if parsing fails
    """
    # Remove markdown code block markers if present
    cleaned_data = raw_data.strip()
    
    # Remove ```json or ```python or other code block markers
    if cleaned_data.startswith('```'):
        cleaned_data = cleaned_data.split('\n', 1)[-1]
    
    # Remove trailing code block marker
    if cleaned_data.endswith('```'):
        cleaned_data = cleaned_data[:-3]
    
    # Remove any leading/trailing whitespace
    cleaned_data = clean_json(cleaned_data).strip()

    
    try:
        # Attempt to parse the cleaned data as JSON
        data_dict = json.loads(cleaned_data)
        return data_dict
    except json.JSONDecodeError as e:
        print(f"Error converting string to dictionary: {e}")
        print(f"Problematic input: {cleaned_data}")
        return False
    
def save_to_word(data, main_key, output_folder, language, outline_data=None):
    if not isinstance(data, dict):
        print("Error: Input data is not a dictionary.")
        return

    doc = Document()

    # Thêm phần meta
    doc.add_paragraph("meta: " + data.get("meta", "")).style = 'Normal'

    # Thêm heading H1
    h1 = doc.add_heading(level=1)
    run = h1.add_run(data.get("h1", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Thêm phần giới thiệu
    doc.add_paragraph(data.get("intro", ""))


    outline_data = outline_data.split("\n") if isinstance(outline_data, str) else outline_data

    # Sử dụng biến outline nếu có (outline là danh sách các dòng với định dạng giống file txt)
    if outline_data != ['']:
        print("Using headings from outline variable")
        for line in outline_data:
            line = line.strip()
            if line.lower().startswith("[h2]"):
                h2 = doc.add_heading(level=2)
                run = h2.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(17)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.lower().startswith("[h3]"):
                h3 = doc.add_heading(level=3)
                run = h3.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph()
    else:
        # Nếu không có outline, sử dụng dữ liệu từ data["h2"]
        for section in data.get("h2", []):
            h2 = doc.add_heading(level=2)
            run = h2.add_run(section.get("title", ""))
            run.bold = True
            run.font.size = Pt(17)
            run.font.color.rgb = RGBColor(0, 0, 0)

            for sub_section in section.get("h3", []):
                if isinstance(sub_section, dict):
                    h3 = doc.add_heading(level=3)
                    run = h3.add_run(sub_section.get("title", ""))
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    doc.add_paragraph()
                elif isinstance(sub_section, str):
                    h3 = doc.add_heading(level=3)
                    run = h3.add_run(sub_section)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    doc.add_paragraph()

    # Dictionary ánh xạ tiêu đề "Kết luận" theo 20 ngôn ngữ phổ biến
    conclusion_titles = {
        "english": "Conclusion",
        "vietnamese": "Kết luận",
        "japanese": "結論",
        "chinese": "结论",
        "korean": "결론",
        "french": "Conclusion",
        "german": "Fazit",
        "spanish": "Conclusión",
        "portuguese": "Conclusão",
        "russian": "Вывод",
        "italian": "Conclusione",
        "dutch": "Conclusie",
        "arabic": "الخاتمة",
        "turkish": "Sonuç",
        "thai": "ข้อสรุป",
        "hindi": "निष्कर्ष",
        "bengali": "উপসংহার",
        "indonesian": "Kesimpulan",
        "malay": "Kesimpulan",
        "polish": "Wniosek"
    }

    

    # Lấy tiêu đề phù hợp, nếu không có thì mặc định dùng tiếng Anh
    conclusion_text = conclusion_titles.get(language.lower(), "Conclusion")

    # Thêm phần kết luận với tiêu đề theo ngôn ngữ
    conclusion_heading = doc.add_heading(level=2)
    run = conclusion_heading.add_run(conclusion_text)
    
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph(data.get("conclusion", ""))

    # Đảm bảo thư mục đầu ra tồn tại
    os.makedirs(output_folder, exist_ok=True)

    # Tạo tên file đầu ra và lưu file
    output_path = os.path.join(output_folder, f"{main_key}.docx")
    doc.save(output_path)
    print(f"File saved to: {output_path}")




# Hàm kiểm tra nếu file đã tồn tại trong thư mục đích
def is_file_processed(output_folder, main_key):
    output_file_path = os.path.join(output_folder, f"{main_key}.docx")
    return os.path.exists(output_file_path)


def generate_seo_outline_api(main_key, secondary_keywords, number_of_h2, number_of_h3, model_name, model_type, api_key, learn_data, language, seo_category):

    prompt = f"""
   
        **LƯU Ý QUAN TRỌNG:** Bài viết phải TUYỆT ĐỐI ĐỘC QUYỀN bằng {language}. KHÔNG SAO CHÉP, KHÔNG DỊCH THUẬT, KHÔNG SỬ DỤNG BẤT KỲ NGÔN NGỮ NÀO KHÁC TRONG BÀI VIẾT.

        **Dữ liệu đầu vào:**
        * Dữ liệu cần học: {learn_data} (Phân tích kỹ lưỡng nếu có để khai thác ý tưởng độc đáo).
        * Thể loại bài SEO: {seo_category}
        * Từ khóa chính: "{main_key}"
        * Từ khóa phụ: "{secondary_keywords}" (Sử dụng từ khóa phụ để mở rộng chủ đề, đào sâu các khía cạnh liên quan, và tạo ra sự đa dạng cho dàn bài).

        **Yêu cầu chi tiết:**
        * **Cấu trúc:** Pyramid HOẶC AIDA.
        * **Ngôn ngữ:** KHÔNG dùng câu hỏi trong tiêu đề, chỉ dùng câu khẳng định tích cực THƯƠNG HIỆU.
        * **Định dạng trả về:** JSON chuẩn (dictionary), không markdown, không index, không "step", không dấu hai chấm sau heading.
        * **Từ khóa:**
            * Từ khóa chính: DUY NHẤT 1 lần trong meta description, intro, conclusion.
            * Từ khóa thương hiệu (hoặc "chúng tôi" nếu không có): DUY NHẤT 1 lần trong meta description, intro, conclusion.

        **Các phần của dàn bài:**

        1. **Meta description (160-180 kí tự):**
            * CHỈ chứa DUY NHẤT 1 từ khóa chính và 1 từ khóa thương hiệu (hoặc "chúng tôi").
            * KHÔNG dùng icon.
            * **Tập trung KHƠI GỢI TỘT ĐỘ SỰ TÒ MÒ, kích thích MẠNH MẼ việc click chuột.**  Viết như một câu "mồi nhử" hấp dẫn nhất.

        2. **Tiêu đề chính (H1) (tối đa 60 kí tự):**
            * CHỨA từ khóa chính.
            * **SÁNG TẠO VƯỢT BẬC, GIẬT TÍT ĐỘC ĐÁO, ĐA DẠNG, HẤP DẪN TUYỆT ĐỐI.**
            * Liên quan đến từ khóa phụ để mở rộng ý tưởng.
            * Thu hút người đọc NGAY LẬP TỨC từ cái nhìn đầu tiên.

        3. **Giới thiệu (Intro) (4-5 dòng, 3-4 câu):**
            * CHỈ chứa DUY NHẤT 1 từ khóa chính.
            * Giải thích từ khóa chính, từ khóa phụ (ngắn gọn).
            * **Khơi gợi sự HỨNG THÚ CAO ĐỘ đối với nội dung chính của bài viết.**  Đặt vấn đề một cách lôi cuốn, khiến người đọc không thể bỏ qua.

        4. **Các tiêu đề phụ (H2) và tiêu đề con (H3):**
            * **KHÁM PHÁ NHỮNG GÓC ĐỘ TIẾP CẬN ĐẦY ĐỦ, THEO ĐÚNG MÔ HÌNH ĐÃ XÁC ĐỊNH Ở TRÊN KHI ĐỘC GIẢ TÌM HIỂU VỀ KEYWORD**
            **KHÔNG DÙNG DẤU ':' hay '-'. TRÁNH VIẾT HAI Ý TRONG MỘT TIÊU ĐỀ, CHỈ TRUYỀN ĐẠT MỘT THÔNG ĐIỆP DUY NHẤT**
            * **TRÁNH TUYỆT ĐỐI LẶP LẠI, KHÔNG ĐI THEO LỐI MÒN.**
            * KHI CHIA SẺ VỀ KINH NGHIỆM CÁ NHÂN, HÃY ĐẶT BẢN THÂN MÌNH VÀO, CHIA SẺ NHIỀU PHƯƠNG PHÁP, GÓC NHÌN SÁNG TẠO. **TUYỆT ĐỐI KHÔNG ĐI THEO LỐI MÒN**
            * Số lượng H2: {number_of_h2} H2.
            * Số lượng H3 mỗi H2: {number_of_h3} đến {number_of_h3+1} H3.
            * **Tiêu đề H3: CỤ THỂ, CHI TIẾT, HỖ TRỢ LÀM RÕ Ý CHO H2, TẠO MẠCH LẠC, LOGIC, NHƯNG VẪN ĐẢM BẢO SỰ SÁNG TẠO VÀ ĐỘC ĐÁO.**
            * **Các tiêu đề H2, H3 phải ĐỘC LẬP, KHÔNG LẶP LẠI Ý TƯỞNG.**

        5. **Kết luận (Conclusion) (4-5 dòng, tối đa 440 kí tự):**
            * CHỈ chứa DUY NHẤT 1 từ khóa chính.
            * Văn phong báo chí tự nhiên, KHÔNG TRỊNH TRỌNG, KHÔ CỨNG.
            * **TÓM TẮT, ĐÚC KẾT CÁC Ý CHÍNH MỘT CÁCH NGẮN GỌN, SÚC TÍCH, ĐỂ LẠI ẤN TƯỢNG SÂU SẮC KHÓ PHAI.**
            * **CALL TO ACTION NẾU CÓ GIỚI THIỆU**
            * **Kết thúc phải THẬT SỰ ĐẮT GIÁ, khiến người đọc suy ngẫm.**

        **TƯ DUY NHƯ NGƯỜI ĐỌC:** Đặt mình vào vị trí người đọc để hiểu RÕ NHẤT NHU CẦU, MONG MUỐN, VẤN ĐỀ của họ liên quan đến "{main_key}". Dàn bài phải GIẢI QUYẾT TRIỆT ĐỂ những điều đó, đồng thời MANG ĐẾN NHỮNG GIÁ TRỊ VƯỢT XA MONG ĐỢI, những ĐIỀU BẤT NGỜ THÚ VỊ. Dẫn dắt người đọc một cách LOGIC, HẤP DẪN, từ cơ bản đến chuyên sâu, nhưng LUÔN GIỮ SỰ SÁNG TẠO VÀ ĐỘC ĐÁO LÀM KIM CHỈ NAM.

        **TỐI ƯU SEO:** Cấu trúc bài viết HỢP LÝ, TỐI ƯU SEO.

        **ĐỊNH DẠNG JSON TRẢ VỀ:**

            {{
            "meta": "<string>: Meta description (160-180 kí tự), chứa duy nhất 1 từ khóa chính và 1 từ khóa thương hiệu hoặc 'chúng tôi'.",
            "h1": "<string>: Tiêu đề chính (H1), tối đa 60 kí tự, chứa từ khóa chính.",
            "intro": "<string>: Phần giới thiệu gồm 3-4 câu, 4-5 dòng, giải thích từ khóa chính, từ khóa phụ và các heading; chỉ chứa duy nhất 1 từ khóa chính.",
            "h2": [
                {{
                "title": "<string>: Tiêu đề phụ (H2), sáng tạo, không lặp lại và không chứa dấu câu đặc biệt như dấu hai chấm hoặc index.",
                "h3": [
                    {{
                    "title": "<string>: Tiêu đề con (H3) cho H2 này, sáng tạo và độc lập."
                    }}
                    // Số lượng H3 tuỳ thuộc vào H2 theo hướng dẫn: H2 ở giữa có {{number_of_h3}} + 1 H3, các H2 còn lại có {{number_of_3}} H3.
                ]
                }}
                // Các mục H2 khác...
            ],
            "conclusion": "<string>: Phần kết luận gồm 3-4 câu, tối đa 440 kí tự, chỉ chứa duy nhất 1 từ khóa chính; tóm tắt và đúc kết các ý chính của bài viết."
            }}
    """

    outline = call_chatbot(prompt, model_name, model_type, api_key) 

        # Trả về dàn bài
    return outline


def generate_outline(keywords, output_folder, number_of_h2, number_of_h3, model_name, model_type,  api_key, learn_data, language, seo_category, outline_data):
    try:
        # Tách chuỗi thành các phần dựa trên dấu phẩy hoặc tab
        if '\t' in keywords:
            parts = keywords.strip().split('\t', 1)  # Tách theo tab, chỉ tách 1 lần
        else:
            parts = keywords.strip().split(',', 1)  # Tách theo dấu phẩy, chỉ tách 1 lần

        # Nếu có ít nhất 1 từ khóa
        if len(parts) == 1:
            main_key = parts[0].strip()  # Chỉ có main_key, không có secondary_keywords
            secondary_keywords = ''
        else:
            main_key = parts[0].strip()  # Phần đầu tiên là main_key
            secondary_keywords = parts[1].strip()  # Phần còn lại là secondary_keywords


        # Kiểm tra nếu file đã tồn tại, nếu có thì bỏ qua
        if is_file_processed(output_folder, main_key):
            return  # Skip this line if file already exists
        
        if learn_data:
            print(f'có learn_data: {learn_data[:10]}')
        else:
        # In ra để kiểm tra
            print(f"main_key: {main_key}")
            


        response = generate_seo_outline_api(main_key, secondary_keywords, number_of_h2, number_of_h3, model_name, model_type, api_key, learn_data, language, seo_category)



        data = convert_to_dict(response)

        if data:
            try:
                save_to_word(data, main_key, output_folder, language, outline_data)  # Lưu dàn bài vào file Word
            except:
                print('Fail to save')


    except Exception as e:
        print(f"An error occurred: {e}")  # In thông báo lỗi
