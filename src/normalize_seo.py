import os
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import string
from src.chat_bot import call_chatbot

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def generate_h1(doc, key_word, model_name, model_type, api_key, language):
    try:
        paragraphs = doc.paragraphs
        
        if len(paragraphs) < 2:
            return "File không có đủ đoạn văn để xử lý."

        # Lấy đoạn văn H1 đầu tiên
        h1_paragraph = paragraphs[1]
        original_text = h1_paragraph.text.strip()


        if model_type == 'chatgpt':
            model_name = 'gpt-4o'

        prompt = f"""
            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**
            Viết lại tiêu đề 1 sau đây để bắt đầu bằng từ khóa '{key_word}', tiêu đề là: {original_text}.
            Keyword đứng đầu theo format: {key_word}: ...
            Đảm bảo viết hoa các chữ cái đầu trong mỗi từ.
            Độ dài từ 45 đến 60 ký tự (bao gồm khoảng trắng). 
            Chỉ gửi lại tiêu đề hoàn chỉnh sau khi sửa, không kèm lời dẫn và không sử dụng định dạng Markdown.
            Viết **SÁNG TẠO**, BAY BỔNG, thu hút người đọc. 
            không gửi kèm lời dẫn, MARKDOWN."""

        while True:
            h1_processed = call_chatbot(prompt, model_name, model_type, api_key)
            if h1_processed.lower().startswith(key_word.lower()) and 45 <= len(h1_processed) <= 60:
                break

        h1_paragraph.text = h1_processed
        h1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h1_paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)

        return doc

    except Exception as e:
        return f"Đã xảy ra lỗi khi xử lý file: {e}"


def generate_h2(doc, key_word, model_name, model_type, api_key, language):
    try:
        h2_paragraphs = [p for p in doc.paragraphs if p.style.name.startswith('Heading 2')]

        if not h2_paragraphs:
            return "File không có tiêu đề H2 để xử lý."

        h2_paragraph = h2_paragraphs[0]
        original_text = h2_paragraph.text.strip()

        if original_text.lower().startswith(key_word.lower()):
            return doc

        if model_type == 'chatgpt':
            model_name = 'gpt-4o'

        prompt = f"""
            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**
            Viết lại tiêu đề H2 sau, đảm bảo 100% có từ khóa '{key_word}' đứng đầu câu.
            Chỉ gửi lại tiêu đề hoàn chỉnh sau khi sửa, không kèm lời dẫn và không sử dụng định dạng Markdown.
            Không thay đổi ý nghĩa nội dung, hãy viết tiêu đề sáng tạo theo từ khóa:\n{original_text}
        """

        while True:
            h2_processed = call_chatbot(prompt, model_name, model_type, api_key)
            if h2_processed.lower().startswith(key_word.lower()):
                break

        h2_paragraph.text = h2_processed
        h2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        return doc

    except Exception as e:
        return f"Đã xảy ra lỗi khi xử lý tiêu đề H2: {e}"


def generate_h3(doc, key_word, model_name, model_type, api_key, language):
    try:
        h3_paragraphs = [p for p in doc.paragraphs if p.style.name.startswith('Heading 3')]

        if not h3_paragraphs:
            return "File không có tiêu đề H3 để xử lý."

        h3_paragraph = h3_paragraphs[0]
        original_text = h3_paragraph.text.strip()

        if original_text.lower().startswith(key_word.lower()):
            return doc

        if model_type == 'chatgpt':
            model_name = 'gpt-4o'

        prompt = f"""
            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**
            Viết lại tiêu đề H3 sau, đảm bảo 100% có từ khóa '{key_word}' đứng đầu câu.
            Chỉ gửi lại tiêu đề hoàn chỉnh sau khi sửa, không kèm lời dẫn và không sử dụng định dạng Markdown.
            Không thay đổi ý nghĩa nội dung, hãy viết tiêu đề sáng tạo theo từ khóa:\n{original_text}
        """

        while True:
            h3_processed = call_chatbot(prompt, model_name, model_type, api_key)
            if h3_processed.lower().startswith(key_word.lower()):
                break

        h3_paragraph.text = h3_processed
        h3_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        return doc

    except Exception as e:
        return f"Đã xảy ra lỗi khi xử lý tiêu đề H3: {e}"


def process_h2_except_first(doc, model_name, model_type, api_key, language):
    try:
        # Tìm tất cả các đoạn văn có kiểu Heading 2
        h2_paragraphs = [p for p in doc.paragraphs if p.style.name.startswith('Heading 2')]

        if len(h2_paragraphs) <= 1:
            return "Không có tiêu đề H2 ngoài tiêu đề đầu tiên để xử lý."

        # Lặp qua các đoạn văn H2 trừ đoạn H2 cuối cùng và xử lý
        for h2_paragraph in h2_paragraphs[:-1]:  # Chỉ lặp qua các H2 trừ H2 cuối
            original_text = h2_paragraph.text.strip()

            # Yêu cầu ChatGPT viết lại tiêu đề H2 với từ khóa 'key_word'
            prompt = f"""
                **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

                Viết lại tiêu đề H2 sau, độ dài tối đa 60 character. 
                Chỉ gửi lại tiêu đề hoàn chỉnh sau khi sửa, không kèm lời dẫn và không sử dụng định dạng Markdown, 
                không kèm nội dung dẫn vào, Viết có nghĩa, không thay đổi nội dung, giữ nguyên các con số:\n{original_text}
            """
            h2_processed = call_chatbot(prompt, model_name, model_type, api_key,)

            # Chỉnh sửa nội dung đoạn văn với tiêu đề đã xử lý
            for run in h2_paragraph.runs:
                run.text = h2_processed

            # Định dạng đoạn văn H2 (ví dụ: căn trái và in nghiêng)
            h2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h2_paragraph.runs:
                run.italic = False  # Bạn có thể thay đổi định dạng tùy ý

        return doc

    except Exception as e:
        return f"Đã xảy ra lỗi khi xử lý tiêu đề H2: {e}"



def process_h3(doc, model_name, model_type, api_key, language):
    try:
        # Duyệt qua toàn bộ các đoạn văn trong tài liệu
        for i, paragraph in enumerate(doc.paragraphs):
            # Kiểm tra xem đoạn văn có kiểu Heading 3 không
            if paragraph.style.name.startswith('Heading 3'):
                # Lấy tiêu đề H3 ban đầu
                original_text = paragraph.text.strip()

                # Lấy văn bản sau tiêu đề H3 (nếu có) để tham khảo
                next_paragraph_text = ""
                if i + 1 < len(doc.paragraphs):
                    next_paragraph = doc.paragraphs[i + 1]
                    next_paragraph_text = next_paragraph.text.strip()

                # Tạo prompt gửi API ChatGPT
                prompt = f"""
                    **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

                    f"Viết lại tiêu đề H3 sau, giữ nguyên ý nghĩa và đảm bảo độ dài tối đa 65 character (bao gồm khoảng trắng).\n
                    f"Chỉ gửi lại tiêu đề hoàn chỉnh sau khi sửa, không kèm lời dẫn và không sử dụng định dạng Markdown, 
                    f"Tiêu đề H3 ban đầu: {original_text}\n\n
                    f"Viết có nghĩa, không thay đổi nội dung, giữ nguyên các con số.
                    f"Văn bản sau H3 (chỉ tham khảo, không viết lại): {next_paragraph_text}\n\n
                """

                # Gọi API ChatGPT để xử lý tiêu đề H3
                h3_processed = call_chatbot(prompt, model_name, model_type, api_key,)
                
                # Cập nhật lại tiêu đề H3 trong tài liệu với tiêu đề đã được sửa
                for run in paragraph.runs:
                    run.text = h3_processed

                # Định dạng lại đoạn văn H3 (căn trái, không in nghiêng)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.italic = False  # Giữ định dạng không in nghiêng

        return doc

    except Exception as e:
        return f"Đã xảy ra lỗi khi xử lý tiêu đề H3: {e}"


def justify_and_set_line_spacing(doc):
    try:
        for i, paragraph in enumerate(doc.paragraphs):
            if i != 1:  # Đảm bảo rằng đoạn văn số 2 (index 1) không bị thay đổi
                # Căn đều đoạn văn
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Cài đặt giãn cách dòng 1.5
                paragraph.paragraph_format.line_spacing = 1.5
                
                # Cài đặt giãn cách đoạn trước và sau
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(10)
        
        return doc

    except Exception as e:
        return f"Lỗi xử lý căn đều và giãn cách dòng: {e}"

def normalize_headings(doc, model_name, model_type, api_key, language):
    try:
        num_paragraphs = len(doc.paragraphs)
        start_index = 3  # Đoạn số 5 từ trên xuống
        end_index = num_paragraphs - 2  # Đoạn số 2 từ dưới lên

        for i in range(start_index, end_index + 1):
            paragraph = doc.paragraphs[i]

            if paragraph.style.name == 'Heading 1':
                paragraph.text = paragraph.text.title()  # Viết hoa chữ cái đầu từng từ
                for run in paragraph.runs:
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen

            elif paragraph.style.name in ['Heading 2', 'Heading 3']:
                original_runs = [(run.text, run.bold, run.font.size) for run in paragraph.runs]

                for run in paragraph.runs:
                    run_text = run.text.lower().capitalize()
                    if run_text and run_text[-1] in string.punctuation:
                        run_text = run_text[:-1]
                    run.text = run_text

                prompt = f""" 
                **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

                Sửa lỗi chính tả cho heading sau: {paragraph.text}.
                Chỉ viết hoa chữ cái đầu tiên của câu, sau dấu ':' và tên riêng.
                **TUYỆT ĐỐI KHÔNG VIẾT HOA TOÀN BỘ CHỮ CÁI ĐẦU MỖI TỪ (TRỪ TÊN RIÊNG)**
                Chỉ gửi lại văn bản đã chuẩn hóa, không kèm lời dẫn, không kèm markdown.
                Đây là heading, không dùng dấu chấm kết thúc câu.
                """

                corrected_text = call_chatbot(prompt, model_name, model_type, api_key)

                corrected_words = corrected_text.split()
                current_index = 0

                for run, (original_text, original_bold, original_size) in zip(paragraph.runs, original_runs):
                    run_words = original_text.split()
                    num_words = len(run_words)

                    run.text = " ".join(corrected_words[current_index:current_index + num_words])
                    run.bold = original_bold
                    run.font.size = Pt(17) if paragraph.style.name == 'Heading 2' else Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen

                    current_index += num_words

        return doc

    except Exception as e:
        return f"Lỗi xử lý file: {e}"




def insert_keyword_below_h3(doc, key_word, max_key, model_name, model_type, api_key, language): 
    """
    Chèn keyword vào đoạn văn ngay dưới các tiêu đề H3 (Heading 3) 
    trong tài liệu Word, áp dụng điều kiện số keyword tối đa.
    
    Args:
        doc: Đối tượng tài liệu Word.
        key_word: Từ khóa cần chèn.
    
    Returns:
        doc: Đối tượng tài liệu Word đã được cập nhật hoặc thông báo lỗi.
    """
    try:

        # Đếm số lần xuất hiện của keyword trong toàn bộ tài liệu (không phân biệt hoa thường)
        total_keyword_count = sum(paragraph.text.lower().count(key_word.lower()) for paragraph in doc.paragraphs)
        print(total_keyword_count)
        # Lưu số lượng keyword hiện tại
        current_keyword_count = total_keyword_count

        num_paragraphs = len(doc.paragraphs)

        # Duyệt tất cả các đoạn văn trong tài liệu
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name == 'Heading 3' and paragraph.runs:
                # Kiểm tra đoạn văn ngay sau H3
                if i + 1 < num_paragraphs:
                    next_paragraph = doc.paragraphs[i + 1]

                    # Kiểm tra nếu đoạn chưa có keyword
                    if key_word.lower() not in next_paragraph.text.lower():
                        # Nếu số lượng keyword đã đủ, dừng chèn thêm
                        if current_keyword_count >= max_key:
                            break

                        # Gọi ChatGPT để cập nhật đoạn văn sau H3
                        prompt = f"""
                                **Lưu ý**: **Tất cả nội dung phải được viết **độc quyền bằng {language}**. 
                                           **Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

                                Đoạn văn gốc: {next_paragraph.text}  
                                Từ khóa: {key_word}  

                                Yêu cầu:  
                                - Chèn thêm 1 từ khóa vào đoạn văn gốc.
                                - Từ khóa phải chèn tự nhiên, mạch lạc, không làm lủng củng câu văn.  
                                - Chèn từ khóa gốc 100%.
                                - Chỉ chỉnh sửa câu chèn từ khóa, các câu khác giữ nguyên, không được chỉnh sửa.  
                                - Từ khóa phải viết thường (trừ khi đứng đầu câu hoặc là tên riêng).  
                                - Viết hoa đúng tên người, địa điểm, hoặc thương hiệu.  

                                Kết quả:  
                                - Trả về đoạn văn hoàn chỉnh sau khi chỉnh sửa, không markdown, không lời dẫn, không giải thích, gửi kết quả tốt nhất.  
                        """
                        update_paragraph = call_chatbot(prompt, model_name, model_type, api_key,)
                        

                        # Giữ nguyên định dạng của đoạn văn
                        next_paragraph.clear()
                        next_paragraph.add_run(update_paragraph)

                        # Cập nhật số lượng keyword sau khi gọi API
                        current_keyword_count = count_keywords_in_document(doc, key_word)

        return doc

    except Exception as e:
        return f"Lỗi xử lý file: {e}"

def insert_keyword(doc, key_word, max_key, model_name, model_type, api_key, language):
    try:
        # Đếm số lần xuất hiện của keyword trong toàn bộ tài liệu (không phân biệt hoa thường)
        total_keyword_count = sum(paragraph.text.lower().count(key_word.lower()) for paragraph in doc.paragraphs)
        
        # Lưu số lượng keyword hiện tại
        current_keyword_count = total_keyword_count

        # Xác định số đoạn cần duyệt từ đoạn 5 từ trên xuống và đoạn 2 từ dưới lên
        num_paragraphs = len(doc.paragraphs)
        start_index = 4  # Đoạn số 5 từ trên xuống (index bắt đầu từ 0)
        end_index = num_paragraphs - 2  # Đoạn số 2 từ dưới lên

        skip = False

        # Duyệt các đoạn văn từ đoạn 5 từ trên xuống và đoạn 2 từ dưới lên
        for i in range(start_index, end_index + 1):
            paragraph = doc.paragraphs[i]
            if paragraph.style.name == 'Heading 3' and paragraph.runs:
                for run in paragraph.runs:
                    # Kiểm tra điều kiện bold và kích thước font
                    if run.bold and run.font.size and run.font.size.pt == 13:
                        # Kiểm tra nếu đoạn chưa có keyword thì gọi ChatGPT
                        if key_word not in paragraph.text:
                            # Nếu số lượng keyword đã đủ (>= 21), dừng gọi API
                            if current_keyword_count >= max_key:
                                return doc

                            prompt = f"""
                            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

                            Viết lại tiêu đề sau với từ khóa '{key_word}':

                            Tiêu đề gốc: {paragraph.text}

                            Yêu cầu:
                            - Chứa chính xác từ khóa một lần, chèn tự nhiên
                            - Tối đa 55 ký tự
                            - Đảm bảo 100% giữ nguyên nội dung chính và các con số, thông tin có trong đoạn, viết tự nhiên.
                            - Chèn từ khóa gốc 100%, không tự đảo vị trí (từ khóa có thể là ghép từ có nghĩa với tên thương hiệu, nên chèn tự nhiên)
                            - Chỉ trả về tiêu đề mới, không kèm định dạng hay giải thích
                            """

                            # Gọi ChatGPT để cập nhật tiêu đề
                            update_h3 = call_chatbot(prompt, model_name, model_type, api_key,)

                            # Loại bỏ ký tự đặc biệt và khoảng trắng thừa
                            update_h3_cleaned = re.sub(r'[^\w\s]', '', update_h3)  # Loại bỏ ký tự đặc biệt
                            update_h3_cleaned = re.sub(r'\s+', ' ', update_h3_cleaned).strip()  # Loại bỏ khoảng trắng thừa

                            # Giữ nguyên định dạng của đoạn văn
                            paragraph.clear()
                            run_updated = paragraph.add_run(update_h3_cleaned)

                            # Sao chép lại các thuộc tính định dạng từ run cũ
                            run_updated.bold = run.bold
                            run_updated.font.size = run.font.size
                            run_updated.font.name = run.font.name  # Đảm bảo font chữ giữ nguyên
                            run_updated.font.color.rgb = run.font.color.rgb  # Đảm bảo màu chữ giữ nguyên

                            # Cập nhật số lượng keyword sau khi gọi API
                            current_keyword_count = count_keywords_in_document(doc, key_word)

                            skip = True
                            break

        return doc

    except Exception as e:
        return f"Lỗi xử lý file: {e}"



def copy_and_insert_h3(doc, key_word, num_h3, model_name, model_type, api_key, language):
    """Chọn num_h3 H3 cách đều nhau trong tài liệu, tạo bản sao và chèn vào trước đoạn H3 tiếp theo hoặc H2 nếu là H3 cuối cùng."""
    # Lấy danh sách các H3 trong tài liệu
    h3_list = [paragraph for paragraph in doc.paragraphs if paragraph.style.name == 'Heading 3']

    if len(h3_list) < num_h3:
        print(f"Tài liệu không có đủ {num_h3} Heading 3 để xử lý.")
        return doc

    # Tính toán khoảng cách giữa các H3 để cách đều
    step = len(h3_list) // (num_h3 - 1) if num_h3 > 1 else len(h3_list)

    # Lấy các H3 cách đều
    selected_h3 = [h3_list[i * step] for i in range(num_h3 - 1)]
    selected_h3.append(h3_list[-1])  # Đảm bảo lấy H3 cuối cùng

    def insert_before_next_h3_or_h2(original_para, text, doc, style='Normal', font_size=Pt(11),
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True):
        """Chèn đoạn văn bản mới trước đoạn H3 tiếp theo hoặc H2 nếu là H3 cuối cùng."""
        # Tạo đoạn văn mới
        new_para = deepcopy(original_para)
        new_para.text = text
        new_para.style = style

        # Thiết lập định dạng font cho từng run trong đoạn văn mới
        for run in new_para.runs:
            run.font.size = font_size
            run.font.italic = italic

        # Thiết lập căn chỉnh đoạn văn mới
        new_para.alignment = alignment

        parent = original_para._element.getparent()
        original_index = list(parent).index(original_para._element)

        # Tìm H3 hoặc H2 tiếp theo
        for next_index in range(original_index + 1, len(doc.paragraphs)):
            next_para = doc.paragraphs[next_index]
            if next_para.style.name in ['Heading 3', 'Heading 2']:
                parent.insert(next_index, new_para._element)
                return new_para

        # Nếu không có H3 hoặc H2 nào sau, chèn vào cuối tài liệu
        parent.append(new_para._element)
        return new_para

    for h3 in selected_h3:
        # Lấy đoạn văn ngay sau h3 (nếu có)
        try:
            h3_index = doc.paragraphs.index(h3)
            # Nếu có đoạn sau h3, dùng đoạn đó, ngược lại dùng chính h3.text
            next_para_text = doc.paragraphs[h3_index + 1].text if h3_index + 1 < len(doc.paragraphs) else h3.text
        except ValueError:
            # Trong trường hợp không tìm thấy (không nên xảy ra) thì dùng h3.text
            next_para_text = h3.text

        prompt = f"""
                    **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**
                    Tiêu đề: {h3.text}. \
                    Từ khóa: {key_word}. \
                    Đoạn SEO: {next_para_text}. \
                    Viết dùm phụ đề ảnh tương ứng với nội dung đoạn SEO và tiêu đề. Tối đa 45 kí tự. Ko '.' kết thúc câu.\
                    Chỉ viết hoa chữ đầu tiên trong câu và tên riêng.\
                    Nếu trong tiêu đề có tồn tại từ khóa, thì phụ đề ảnh 100% có từ khóa.\
                    Nếu trong tiêu đề không có từ khóa, thì phụ đề ảnh đảm bảo không có từ khóa.\
                    Chỉ gửi về phụ đề ảnh. Ko mark down, ko lời dẫn.\
                    Gửi phiên bản tốt nhất, không đưa ra lựa chọn. \
                    """

        response = call_chatbot(prompt, model_name, model_type, api_key,)

        # Chèn đoạn văn bản mới trước H3 tiếp theo hoặc H2 nếu là H3 cuối cùng
        insert_before_next_h3_or_h2(
            original_para=h3,
            text=response,
            doc=doc,
            style='Normal',
            font_size=Pt(11),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True
        )

    return doc


def insert_paragraph_after(paragraph, text="", style=None, font_size=None, bold=False, italic=False, alignment=None):
    """
    Chèn một đoạn văn mới sau đoạn văn hiện tại.

    :param paragraph: Đoạn văn hiện tại (Paragraph) sau đó sẽ chèn đoạn văn mới.
    :param text: Nội dung của đoạn văn mới.
    :param style: Style của đoạn văn mới (str).
    :param font_size: Kích thước font (Pt).
    :param bold: Đặt in đậm (bool).
    :param italic: Đặt in nghiêng (bool).
    :param alignment: Căn chỉnh đoạn văn mới (WD_PARAGRAPH_ALIGNMENT).
    :return: Đoạn văn mới được chèn (Paragraph).
    """
    # Tạo một phần tử đoạn văn mới
    new_p = OxmlElement('w:p')
    # Chèn sau đoạn văn hiện tại
    paragraph._element.addnext(new_p)
    # Tạo đối tượng Paragraph mới từ phần tử XML
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    run = new_paragraph.add_run(text)
    if font_size:
        run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if alignment:
        new_paragraph.alignment = alignment
    return new_paragraph


def fix_text_between_h2_h3(doc, model_name, model_type, api_key, key_word, language):
    """
    Xử lý tài liệu Word để chèn văn bản giữa H2 và H3 nếu sau H2 là H3.
    """
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if para.style.name == 'Heading 2':
            h2_para = para
            if i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]
                if next_para.style.name == 'Heading 3':
                    h2_text = h2_para.text.strip()
                    
                    prompt = f"""
                        Viết 2 câu giới thiệu mở đầu cho đoạn nội dung thuộc heading: "{h2_text}".  
                        Nội dung phải viết **hoàn toàn bằng {language}**, không chứa bất kỳ ngôn ngữ khác.  
                        Không lặp lại "{h2_text}" trong câu, không chèn "{key_word}".  
                        Không sử dụng từ "heading" hoặc đề cập đến "bài viết này".  
                        Chỉ trả về 2 câu trong cùng một đoạn, không có tiêu đề hay markdown.  
                        Tối đa 250 ký tự.  

                    """


                    
                    inserted_text = call_chatbot(prompt, model_name, model_type, api_key)
                                        
                    if inserted_text:
                        inserted_text = inserted_text.strip()
                        if len(inserted_text) > 0:
                            inserted_text = inserted_text[0].upper() + inserted_text[1:]

                        new_para = insert_paragraph_after(
                            paragraph=h2_para,
                            text=inserted_text,
                            style='Normal',
                            font_size=Pt(11),
                            bold=False,
                            italic=False,
                            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        )
                        i += 1
        i += 1
    return doc



def count_characters(text):
    """
    Đếm số ký tự (bao gồm khoảng trắng) trong một đoạn văn bản.
    
    :param text: Đoạn văn bản cần đếm ký tự
    :return: Số ký tự trong đoạn văn bản
    """
    return len(text)

def fix_meta(doc, keyword, model_name, model_type, api_key, language):
    """
    Xử lý đoạn meta (đoạn đầu tiên) trong tài liệu Word bằng cách gửi nội dung đoạn đó cùng với từ khóa tới ChatGPT để viết lại.

    :param doc: Đối tượng Document đã được mở bằng python-docx
    :param keyword: Từ khóa để bao gồm trong đoạn meta mới
    :return: Đối tượng Document sau khi chỉnh sửa
    """
    if not doc.paragraphs:
        print("Tài liệu không có đoạn văn nào để xử lý.")
        return doc


    # Lấy đoạn meta là đoạn đầu tiên
    meta_para = doc.paragraphs[0]
    meta_text = meta_para.text.strip()
    keyword = keyword.lower()
    if not meta_text:
        print("Đoạn meta (đoạn đầu tiên) trống. Không có gì để xử lý.")
        return doc

    # Đếm số ký tự trong đoạn meta hiện tại
    char_count = count_characters(meta_text)

    # Kiểm tra xem số ký tự có thỏa mãn yêu cầu không
    if 125 <= char_count - 7 <= 155 and keyword in meta_text:
        print("Đoạn meta đã thỏa mãn yêu cầu về số ký tự. Bỏ qua việc gọi ChatGPT.")
        return doc


    if model_type  == "chatgpt":
        prompt = f"""
            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

            Hãy viết lại đoạn meta sau đây, đảm bảo 100% có duy nhất 1 từ khóa chính: '{keyword}'. Từ khóa đứng đầu câu chuẩn SEO.
            Đoạn meta nên rõ ràng, súc tích và thu hút người đọc. Độ dài từ 130 đến 180 kí tự tính cả khoảng trắng.
            Đoạn meta hiện tại:\n\n{meta_text}\n\nVăn bản trả về theo cấu trúc sau: 'meta: nội dung'.
        """
    else:
        prompt = f"""
            **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

            Hãy viết lại đoạn meta sau đây, đảm bảo 100% có duy nhất 1 từ khóa chính: '{keyword}'. 
            Đoạn meta nên rõ ràng, súc tích và thu hút người đọc. Độ dài từ 100 đến 160 kí tự tính cả khoảng trắng.
            Đoạn meta hiện tại:\n\n{meta_text}\n\nVăn bản trả về theo cấu trúc sau: 'meta: nội dung'.
        """

    new_meta_text = ""
    # Kiểm tra và yêu cầu ChatGPT trả về đúng độ dài ký tự
    while True:
        # Gọi ChatGPT để nhận đoạn meta mới
        new_meta_text = call_chatbot(prompt, model_name, model_type, api_key,)

        if not new_meta_text:
            print("Không thể nhận được đoạn meta mới từ ChatGPT.")
            return doc
        
        # Đảm bảo rằng đoạn meta mới không chứa các ký tự đặc biệt không cần thiết
        new_meta_text = new_meta_text.replace('\n', ' ').strip()

        # Đếm số ký tự trong đoạn meta mới
        char_count = count_characters(new_meta_text)

        # Kiểm tra số ký tự có đúng yêu cầu (130 <= char_count <= 155)
        if 140 <= char_count <= 155:
            break
        else:
            print(f"Đoạn meta trả về có {char_count} ký tự, không thỏa mãn yêu cầu. Đang thử lại...")

    # Xác định định dạng hiện tại của đoạn meta
    # Giả định rằng tất cả các run trong đoạn meta có cùng định dạng
    runs = meta_para.runs
    if runs:
        first_run = runs[0]
        font_size = first_run.font.size
        bold = first_run.font.bold
        italic = first_run.font.italic
        style = meta_para.style
        alignment = meta_para.alignment
    else:
        # Nếu không có run nào, sử dụng các giá trị mặc định
        font_size = Pt(12)
        bold = False
        italic = False
        style = None
        alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Xóa tất cả các run hiện tại trong đoạn meta
    for run in meta_para.runs:
        run.text = ""

    # Thêm đoạn meta mới vào đoạn meta hiện tại
    new_run = meta_para.add_run(new_meta_text)
    if font_size:
        new_run.font.size = font_size
    new_run.bold = bold
    new_run.italic = italic
    meta_para.style = style
    meta_para.alignment = alignment

    return doc

def count_keywords_in_document(doc, key_word):
    """
    Hàm này sẽ trả về tổng số lần xuất hiện của từ khóa trong tài liệu.
    """
    total_count = 0
    # Duyệt qua tất cả các paragraph trong tài liệu
    for paragraph in doc.paragraphs:
        # Đếm số lần xuất hiện của key_word trong mỗi paragraph (không phân biệt chữ hoa chữ thường)
        total_count += paragraph.text.lower().count(key_word.lower().strip())
    return total_count


def key_reduction(doc, key_word, max_key, model_name, model_type, api_key, language):
    """
    Hàm này giảm số lần xuất hiện của key_word trong tài liệu nếu số lần xuất hiện vượt quá max_key.
    """
    doc = merge_sections(doc)    

    # Đếm tổng số lần xuất hiện của key_word trong toàn bộ tài liệu
    total_count = count_keywords_in_document(doc, key_word)

    # Duyệt qua các paragraph từ paragraph thứ 9 tới paragraph thứ 2 từ dưới lên với bước nhảy là 3
    total_paragraphs = len(doc.paragraphs)
    start_idx = 4  # Tính toán vị trí bắt đầu
    end_idx = total_paragraphs - 2  # Paragraph thứ 2 từ dưới lên
    api_call_count = 0  # Biến đếm số lần gọi API

    # Kiểm tra nếu số lần xuất hiện của key_word > max_key thì tiến hành xử lý
    while total_count > max_key:
        for i in range(start_idx, end_idx, 3):
            paragraph = doc.paragraphs[i]

            # Kiểm tra nếu đoạn văn có định dạng là Heading 2 thì bỏ qua
            if paragraph.style.name == "Heading 2":
                continue

            # Đếm số lần xuất hiện của từ khóa trong paragraph (không phân biệt chữ hoa chữ thường)
            count = paragraph.text.lower().count(key_word.lower())

            # Nếu số lượng từ khóa vượt quá 0 trong paragraph, tiến hành xử lý
            if count > 0 and paragraph.text.lower() != "kết luận":
                # Gọi hàm API (call_chatgpt) để viết lại đoạn văn bản
                prompt = f"""
                **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**
                **TUYỆT ĐỐI GIỮ NGUYÊN MARKDOWN có sẵn: <<<PARA_START:Normal>>>, <<<PARA_START:List Bullet>>>**

                Viết lại đoạn văn sau (có thể là heading title) bằng cách giảm số lần xuất hiện của từ khóa '{key_word}':

                {paragraph.text}

                Yêu cầu:
                - Xóa bớt 1 từ khóa (nếu có nhiều hơn 1), hoặc xóa hoàn toàn (nếu chỉ có 1).
                - Đảm bảo 100% giữ nguyên nội dung chính và các con số cụ thể, thông tin có trong đoạn, viết tự nhiên.
                - Chỉ thay đổi câu chứa từ khóa, câu không chứa từ khóa giữ nguyên 100%.  
                - Chỉ trả về đoạn văn mới, không kèm định dạng hay giải thích hay đưa ra lựa chọn, gửi bản tốt nhất.
                - Chú ý đảm bảo viết hoa tên riêng người, địa điểm.  
                """
                api_call_count += 1
                new_text = call_chatbot(prompt, model_name, model_type, api_key,)

                # Lưu lại tất cả các runs gốc trong paragraph
                runs = paragraph.runs

                # Xóa toàn bộ nội dung cũ trong paragraph
                paragraph.clear()

                # Thêm đoạn văn bản mới
                run = paragraph.add_run(new_text)

                # Sao chép lại tất cả các thuộc tính định dạng từ các run cũ vào run mới
                for orig_run in runs:
                    run.font.bold = orig_run.bold
                    run.font.italic = orig_run.italic
                    run.font.underline = orig_run.underline
                    run.font.size = orig_run.font.size
                    run.font.name = orig_run.font.name
                    run.font.color.rgb = orig_run.font.color.rgb
                    # Thêm các thuộc tính định dạng khác nếu cần

            # Kiểm tra nếu số lần xuất hiện từ khóa đã giảm xuống <= max_key
            if count_keywords_in_document(doc, key_word) <= max_key:
                break

        # Cập nhật lại start_idx và end_idx sau mỗi lần xử lý
        print('Kết thúc vòng lặp, số lần gọi API: ', api_call_count)
        start_idx += 1

        if api_call_count >= 30:
            break
        # Kiểm tra lại tổng số từ khóa trong tài liệu
        total_count = count_keywords_in_document(doc, key_word)
        print("Tổng số từ khóa còn lại:", total_count)


    doc = split_sections(doc)    

    return doc


import re
from docx.shared import RGBColor

def clear_paragraph(paragraph):
    """
    Xóa tất cả các runs trong đoạn văn.
    """
    p = paragraph._element
    for child in p[:]:
        p.remove(child)

def format_keyword_in_paragraph(paragraph, keyword, bold=True, color=RGBColor(255, 0, 0)):
    """
    Định dạng tất cả các xuất hiện của từ khóa trong đoạn văn bằng cách làm in đậm và đổi màu đỏ.
    """
    # Lấy toàn bộ văn bản của đoạn văn
    full_text = paragraph.text
    # Tìm tất cả các vị trí xuất hiện của từ khóa (không phân biệt chữ hoa chữ thường)
    matches = list(re.finditer(re.escape(keyword), full_text, re.IGNORECASE))
    
    if not matches:
        return
    
    # Xóa tất cả các runs hiện tại
    clear_paragraph(paragraph)
    
    last_index = 0
    for match in matches:
        start, end = match.span()
        # Thêm văn bản trước từ khóa
        if start > last_index:
            before_text = full_text[last_index:start]
            run = paragraph.add_run(before_text)
            # Bạn có thể đặt định dạng cho run này nếu cần
        # Thêm từ khóa với định dạng đặc biệt
        keyword_text = full_text[start:end]
        keyword_run = paragraph.add_run(keyword_text)
        keyword_run.bold = bold
        last_index = end
    # Thêm văn bản còn lại sau từ khóa cuối cùng
    if last_index < len(full_text):
        after_text = full_text[last_index:]
        run = paragraph.add_run(after_text)
        # Bạn có thể đặt định dạng cho run này nếu cần

def get_last_non_empty_paragraph(doc):
    """
    Lấy đoạn văn cuối cùng có chứa văn bản (loại trừ các đoạn rỗng hoặc chỉ chứa khoảng trắng).
    """
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():  # Kiểm tra nếu đoạn văn không rỗng sau khi loại bỏ khoảng trắng
            return paragraph
    return None

def split_into_sentences(text):
    """
    Tách đoạn văn thành các câu sử dụng regex.
    """
    # Sử dụng regex để tách câu dựa trên dấu chấm, dấu hỏi hoặc dấu chấm than
    sentence_endings = re.compile(r'(?<=[.!?]) +')
    sentences = sentence_endings.split(text.strip())
    return sentences

def check_intro(doc, key_word, model_name, model_type, api_key, language):
    """
    Kiểm tra và định dạng từ khóa trong đoạn văn thứ 3 của tài liệu.
    - Nếu từ khóa xuất hiện duy nhất 1 lần và nằm trong câu đầu tiên, định dạng từ khóa.
    - Nếu không, gọi ChatGPT để viết lại đoạn văn chứa đúng 1 từ khóa trong câu đầu tiên.
    - Định dạng giãn cách dòng 1.5 và căn đều cho đoạn văn.
    - Xóa khoảng trắng đầu và cuối đoạn văn.
    """
    max_attempts = 3  # Số lần thử tối đa
    attempt = 0
    success = False

    while attempt < max_attempts and not success:
        attempt += 1

        # Kiểm tra xem tài liệu có ít nhất 3 đoạn văn không
        if len(doc.paragraphs) < 3:
            break  # Nếu không có 3 đoạn văn, thoát vòng lặp và trả về tài liệu không thay đổi

        # Lấy đoạn văn thứ 3 (index 2)
        paragraph = doc.paragraphs[2]
        paragraph.text = paragraph.text.strip()  # Xóa khoảng trắng đầu và cuối
        key_word_lower = key_word.lower()
        

        prompt = f"""
        **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

        Đoạn cần viết lại: {paragraph.text}.
        Viết lại đoạn văn giới thiệu cho bài viết này, phong cách nhà báo chuyên nghiệp, ko thay đổi nội dung 
        đảm bảo 100% chứa duy nhất 1 từ khóa "{key_word}" nằm trong câu đầu tiên. 
        Đoạn văn nên rõ ràng, súc tích và phù hợp với nội dung chung của bài viết, 
        Viết có nghĩa, đảm bảo 100% không thay đổi nội dung, giữ nguyên các con số, tên thương hiệu đặc biệt.
        """
        new_text = call_chatbot(prompt, model_name, model_type, api_key,)
        if new_text:
            new_text = new_text.strip()  # Xóa khoảng trắng đầu và cuối
            new_text_lower = new_text.lower()
            new_count = new_text_lower.count(key_word_lower)
            # Tách đoạn văn mới thành các câu để kiểm tra từ khóa trong câu đầu tiên
            new_sentences = split_into_sentences(new_text)
            if new_sentences:
                new_first_sentence = new_sentences[0].lower()
                new_keyword_in_first_sentence = key_word_lower in new_first_sentence
            else:
                new_keyword_in_first_sentence = False

            if new_count == 1 and new_keyword_in_first_sentence:
                paragraph.text = new_text
                format_keyword_in_paragraph(paragraph, key_word)
            else:
                # Nếu ChatGPT không chèn được từ khóa duy nhất trong câu đầu tiên, không thực hiện thay đổi
                continue  # Thử lại
        else:
            continue  # Thử lại nếu không nhận được phản hồi từ ChatGPT

        # Kiểm tra điều kiện có ít nhất 1 keyword được in đậm
        success = False
        for run in paragraph.runs:
            if run.bold and re.search(re.escape(key_word), run.text, re.IGNORECASE):
                success = True
                break

        # Định dạng giãn cách dòng và căn đều cho đoạn văn
        paragraph.paragraph_format.line_spacing = Pt(18)  # Giãn cách dòng 1.5 (18pt)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Căn đều

    return doc

def check_conclusion(doc, key_word, model_name, model_type, api_key, language):
    """
    Kiểm tra và định dạng từ khóa trong đoạn văn kết luận của tài liệu.
    - Nếu từ khóa xuất hiện duy nhất 1 lần, định dạng từ khóa.
    - Nếu không (0 hoặc >= 2), gọi ChatGPT để viết lại đoạn văn chứa đúng 1 từ khóa.
    - Định dạng giãn cách dòng 1.5 và căn đều cho đoạn văn.
    - Xóa khoảng trắng đầu và cuối đoạn văn.
    """
    max_attempts = 3  # Số lần thử tối đa
    attempt = 0
    success = False

    while attempt < max_attempts and not success:
        attempt += 1

        # Lấy đoạn văn kết luận (đoạn văn cuối cùng không rỗng)
        paragraph = get_last_non_empty_paragraph(doc)
        if not paragraph:
            break  # Nếu không có đoạn văn kết luận, thoát vòng lặp và trả về tài liệu không thay đổi

        paragraph.text = paragraph.text.strip()  # Xóa khoảng trắng đầu và cuối
        key_word_lower = key_word.lower()

        # Gọi ChatGPT để viết lại đoạn văn kết luận
        prompt = f"""
        **Lưu ý: Tất cả nội dung phải được viết **độc quyền bằng {language}**. Không được sử dụng bất kỳ từ ngữ, câu hay mô tả nào thuộc ngôn ngữ khác.**

        Đoạn cần viết lại: {paragraph.text}.
        Viết lại đoạn văn kết luận cho bài viết này, phong cách nhà báo chuyên nghiệp, đảm bảo 100% chứa duy nhất 1 khóa "{key_word}". 
        Đoạn văn nên rõ ràng, súc tích và phù hợp với nội dung chung của bài viết, , ko thay đổi nội dung,
        không được bắt đầu bằng "Kết luận".
        Cả đoạn dài tối đa 380 character.
        Viết có nghĩa, đảm bảo 100% không thay đổi nội dung, giữ nguyên các con số, tên thương hiệu đặc biệt.
        """
        
        new_text = call_chatbot(prompt, model_name, model_type, api_key,)
        if new_text and new_text.lower().count(key_word_lower) == 1:
            paragraph.text = new_text.strip()  # Xóa khoảng trắng đầu và cuối
            format_keyword_in_paragraph(paragraph, key_word)
        else:
            # Nếu ChatGPT không chèn được từ khóa duy nhất, không thực hiện thay đổi
            continue  # Thử lại

        # Kiểm tra điều kiện có ít nhất 1 keyword được in đậm
        success = False
        for run in paragraph.runs:
            if run.bold and re.search(re.escape(key_word), run.text, re.IGNORECASE):
                success = True
                break

        # Định dạng giãn cách dòng và căn đều cho đoạn văn
        paragraph.paragraph_format.line_spacing = Pt(18)  # Giãn cách dòng 1.5 (18pt)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Căn đều

    return doc



def change_font_to_arial(doc):
    """
    Changes the font of all text in a docx document to Arial, preserving other formatting.

    Args:
        doc: A docx.Document object (the opened docx file).

    Returns:
        A docx.Document object with the font changed to Arial.  Returns None if an error occurs.
    """
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Preserve font size and other formatting
                original_size = run.font.size
                original_bold = run.font.bold
                original_italic = run.font.italic
                original_underline = run.font.underline
                original_color = run.font.color.rgb

                run.font.name = 'Arial'  #Set font to Arial

                # Restore other formatting attributes
                if original_size:
                    run.font.size = original_size
                if original_bold:
                    run.font.bold = True
                if original_italic:
                    run.font.italic = True
                if original_underline:
                    run.font.underline = original_underline
                if original_color:
                    run.font.color.rgb = original_color


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Preserve font size and other formatting (same as above)
                            original_size = run.font.size
                            original_bold = run.font.bold
                            original_italic = run.font.italic
                            original_underline = run.font.underline
                            original_color = run.font.color.rgb

                            run.font.name = 'Arial' #Set font to Arial

                            # Restore other formatting attributes
                            if original_size:
                                run.font.size = original_size
                            if original_bold:
                                run.font.bold = True
                            if original_italic:
                                run.font.italic = True
                            if original_underline:
                                run.font.underline = original_underline
                            if original_color:
                                run.font.color.rgb = original_color

        return doc
    except Exception as e:
        print(f"An error occurred: {e}")
        return None


def check_images_in_document(doc):
    """Kiểm tra sự tồn tại của tên ảnh và in ra dòng số của đoạn chứa tên ảnh."""
    image_positions = []  # Lưu các vị trí của các ảnh trong tài liệu
    image_names = []      # Lưu tên ảnh (tên đoạn văn chứa tên ảnh)
    
    for idx, para in enumerate(doc.paragraphs):
        # Kiểm tra nếu đoạn văn là căn giữa
        if para.alignment == 1 or 1:  # 1 là mã của alignment=center
            # Kiểm tra nếu đoạn văn có chứa run in nghiêng và đoạn văn không rỗng
            for run in para.runs:
                if run.font.italic and para.text.strip():  # Kiểm tra in nghiêng và đoạn văn không rỗng
                    image_positions.append(idx + 1)  # Lưu dòng số chứa ảnh (dòng bắt đầu từ 1)
                    image_names.append(para.text.strip())  # Lưu tên ảnh (văn bản của đoạn)
                    break  # Thoát khỏi vòng lặp sau khi tìm thấy lần đầu tiên

    return image_positions, image_names


def format_image_names_in_document(doc):
    """Tìm và định dạng tên ảnh trong tài liệu."""
    # Lấy các vị trí và tên ảnh từ hàm check_images_in_document
    image_positions, image_names = check_images_in_document(doc)
    
    for idx, para in zip(image_positions, image_names):
        # Tìm đoạn văn theo vị trí
        paragraph = doc.paragraphs[idx - 1]  # Vị trí trong doc.paragraphs bắt đầu từ 0, nên trừ 1
        
        # Chuyển văn bản thành viết thường, chỉ viết hoa chữ cái đầu
        formatted_text = para.capitalize()  # Sử dụng capitalize để viết hoa chữ đầu
        
        # Cập nhật lại văn bản của đoạn
        paragraph.text = formatted_text
        
        # Định dạng đoạn văn:
        paragraph.alignment = 1  # Căn giữa (center alignment)
        
        # Định dạng các run trong đoạn văn
        for run in paragraph.runs:
            run.font.italic = True  # In nghiêng
            run.font.size = Pt(11)   # Cỡ chữ 11
    
    return doc

def word_reduction(doc, max_word):
    # Đếm tổng số từ trong tài liệu
    def count_words_in_paragraph(paragraph):
        return len(paragraph.text.split())

    # Tính tổng số từ trong tài liệu
    totals_word = sum(count_words_in_paragraph(paragraph) for paragraph in doc.paragraphs)

    heading_count = 0  # Đếm số lượng Heading 3 đã gặp
    i = 0
    while totals_word > max_word:
        if i < len(doc.paragraphs):
            # Nếu gặp Heading 3
            if doc.paragraphs[i].style.name == 'Heading 3':
                heading_count += 1
                # Bắt đầu xử lý từ Heading 3 thứ 3 trở đi
                if heading_count >= 2:
                    if i + 2 < len(doc.paragraphs):
                        # Kiểm tra đoạn thứ 2 sau Heading 3 (i + 2)
                        if doc.paragraphs[i + 2].style.name == 'Normal' and count_words_in_paragraph(doc.paragraphs[i+2]) >= 40:
                            # Xóa hoàn toàn đoạn văn bản thứ 2 sau Heading 3
                            doc.paragraphs[i + 2]._element.getparent().remove(doc.paragraphs[i + 2]._element)
        i += 1
        totals_word = sum(count_words_in_paragraph(paragraph) for paragraph in doc.paragraphs)

        # Dừng nếu đã duyệt hết tài liệu
        if i >= len(doc.paragraphs):
            break

    return doc

def normalize_spaces_in_docx(doc):
    # Lấy tất cả các đoạn văn trong tài liệu
    paragraphs = doc.paragraphs
    num_paragraphs = len(paragraphs)

    # Xử lý toàn bộ đoạn văn, trừ index 2 và -1
    for i, para in enumerate(paragraphs):
        if i == 2 or i == num_paragraphs - 1:
            continue  # Bỏ qua các đoạn văn có index 2 và -1
        for run in para.runs:
            original_text = run.text
            run.text = ' '.join(original_text.split())

    # Trả về tài liệu đã được xử lý
    return doc


def convert_to_list_format(doc):
    # Duyệt qua các paragraph trong tài liệu
    for para in doc.paragraphs:
        # Kiểm tra nếu paragraph bắt đầu bằng dấu "*"
        if para.text.strip().startswith('*'):
            # Loại bỏ dấu '*' trước khi chuyển thành danh sách
            para.text = para.text.lstrip('*').lstrip()  # Xóa dấu * và khoảng trắng sau dấu *
            
            # Thay đổi style của paragraph thành danh sách
            para.style = 'List Bullet'  # Sử dụng style danh sách có gạch đầu dòng
    return doc

def capitalize_heading(text):
    """
    Viết hoa chữ cái đầu của mỗi từ trong chuỗi.
    """
    return ' '.join(word.capitalize() for word in text.split())



def format_run(run, font_name="Times New Roman", font_size=None):
    """Định dạng một đoạn văn bản với font và giữ nguyên các thuộc tính cũ."""
    original_bold = run.font.bold
    original_italic = run.font.italic
    original_underline = run.font.underline
    original_color = run.font.color.rgb
    
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if original_bold:
        run.font.bold = True
    if original_italic:
        run.font.italic = True
    if original_underline:
        run.font.underline = original_underline
    if original_color:
        run.font.color.rgb = original_color


def find_paragraph_with_max_occurrence_in_doc(word, doc, exclude_indices=set()):
    """
    Tìm đoạn văn trong doc có số lần xuất hiện của từ 'word' nhiều nhất,
    bỏ qua các chỉ số trong exclude_indices.
    Trả về tuple (index, paragraph). Nếu không tìm thấy, trả về (None, None).
    """
    max_count = 0
    idx_max = None
    for idx, para in enumerate(doc.paragraphs):
        if idx in exclude_indices:
            continue
        text = para.text
        if not text.strip():
            continue
        count = len(re.findall(r'\b' + re.escape(word) + r'\b', text.lower()))
        if count > max_count:
            max_count = count
            idx_max = idx
    if idx_max is None:
        return None, None
    return idx_max, doc.paragraphs[idx_max]

def get_overall_word_density_in_doc(word, doc):
    """
    Tính mật độ của từ 'word' trong toàn bộ document.
    """
    texts = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    combined_text = " ".join(texts).lower()
    words = re.findall(r'\b\w+\b', combined_text)
    total_words = len(words)
    count = words.count(word.lower())
    return (count / total_words) * 100 if total_words > 0 else 0


def delete_paragraph(paragraph):
    """
    Xóa đoạn văn khỏi tài liệu bằng cách thao tác trực tiếp với XML.
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def insert_paragraph_after(paragraph, text):
    """
    Chèn một đoạn văn mới sau đoạn văn đã cho.
    """
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para

def merge_sections(doc):
    """
    Hàm merge_sections nhận vào một đối tượng Document (đã mở bởi python-docx)
    và gộp toàn bộ nội dung giữa 2 heading (Heading 2 hoặc Heading 3) thành 1 đoạn văn duy nhất,
    không xuống dòng. Mỗi đoạn gốc được đánh dấu bằng marker có dạng:
       <<<PARA_START:{style_name}>>>{original_text}
    Marker này giúp hàm split_sections có thể tái tạo lại nội dung ban đầu.
    """
    heading_styles = {"Heading 2", "Heading 3"}
    paragraphs = doc.paragraphs

    # Lấy danh sách các chỉ số của các heading
    heading_indices = [i for i, p in enumerate(paragraphs) if p.style.name in heading_styles]

    # Duyệt các đoạn giữa các heading theo thứ tự ngược (để không làm thay đổi index khi xóa)
    for seg in range(len(heading_indices) - 1, 0, -1):
        start_idx = heading_indices[seg - 1] + 1
        end_idx = heading_indices[seg]
        if start_idx < end_idx:
            merged_text = ""
            # Duyệt qua các đoạn văn trong segment và tạo marker cho từng đoạn
            for i in range(start_idx, end_idx):
                p = paragraphs[i]
                style_name = p.style.name
                marker = f"<<<PARA_START:{style_name}>>>"
                merged_text += marker + p.text
            # Gán merged_text cho đoạn văn đầu tiên của segment
            paragraphs[start_idx].text = merged_text
            # Xóa các đoạn văn còn lại trong segment
            for i in range(start_idx + 1, end_idx):
                delete_paragraph(paragraphs[i])
    return doc

def split_sections(doc):
    """
    Hàm split_sections nhận vào Document đã được xử lý bởi merge_sections và
    phục hồi lại cấu trúc ban đầu bằng cách sử dụng marker đã lưu.
    
    Cụ thể, với mỗi đoạn văn chứa marker dạng:
       <<<PARA_START:{style_name}>>>{text}
    Hàm sẽ:
       - Lấy ra danh sách các tuple (style, text)
       - Thay thế đoạn hiện tại bằng tuple đầu tiên
       - Với các tuple còn lại, chèn đoạn mới sau đó với style tương ứng
    """
    # Duyệt qua tất cả các đoạn văn trong tài liệu
    for p in doc.paragraphs:
        if "<<<PARA_START:" in p.text:
            # Sử dụng regex để tách marker và nội dung.
            # Pattern: <<<PARA_START:(.*?)>>>(.*?)(?=<<<PARA_START:|$)
            segments = re.findall(r"<<<PARA_START:(.*?)>>>(.*?)(?=<<<PARA_START:|$)", p.text)
            if segments:
                # Xử lý segment đầu tiên: thay thế nội dung và style của đoạn hiện tại
                first_style, first_text = segments[0]
                p.text = first_text
                p.style = first_style
                current_paragraph = p
                # Với các segment còn lại, chèn đoạn sau đoạn hiện tại
                for style, text in segments[1:]:
                    new_para = insert_paragraph_after(current_paragraph, text)
                    new_para.style = style
                    current_paragraph = new_para
    return doc


def process_word_density_doc(doc, keyword, model_name, model_type, api_key, threshold):
    """
    Xử lý document để giảm mật độ các từ có tần suất vượt ngưỡng threshold (mặc định 3%).
    Biến keyword (tên file) được gửi kèm qua prompt để đảm bảo giữ nguyên keyword đó.
    
    Đầu vào:
        - doc: đối tượng Document (đã được mở từ file Word)
        - keyword: chuỗi từ khóa (tên file) cần giữ nguyên trong nội dung
        - threshold: ngưỡng mật độ từ (mặc định là 3%)
    
    Trả về:
        - document đã được cập nhật
    """
    doc = merge_sections(doc)    

    # Hàm kiểm tra xem đoạn văn có phải là heading 1,2,3 hay không
    def is_heading(para):
        return para.style.name in ["Heading 1", "Heading 2", "Heading 3"]
    
    # Tính mật độ từ ban đầu trong document (chỉ tính các đoạn không phải heading)
    all_text = " ".join([para.text for para in doc.paragraphs 
                         if para.text.strip() != "" and not is_heading(para)])
    word_matches = re.findall(r'\b\w+\b', all_text.lower())
    total_words = len(word_matches)
    word_counts = {}
    for word in word_matches:
        word_counts[word] = word_counts.get(word, 0) + 1
    word_density = {word: (count / total_words) * 100 for word, count in word_counts.items()}
    
    unwanted_keys = {"para_start", "normal", "list", "bullet"}

    # Lọc dictionary, kiểm tra key sau khi chuyển về chữ thường
    word_density = {
        key: value for key, value in word_density.items()
        if key.lower() not in unwanted_keys
    }
    # Lấy danh sách các từ có mật độ vượt quá threshold
    words_over_threshold = [word for word, dens in word_density.items() if dens > threshold]
        

    attempts = {}       # Theo dõi số lần gọi API cho từng đoạn (theo chỉ số)
    exclusion_set = set()  # Các đoạn đã gọi API tối đa 5 lần sẽ bị loại khỏi xử lý
    
    # Xử lý từng từ vượt ngưỡng
    for word in words_over_threshold:
        current_density = get_overall_word_density_in_doc(word, doc)
        while current_density >= threshold:
            idx, para = find_paragraph_with_max_occurrence_in_doc(word, doc, exclude_indices=exclusion_set)
            if idx is None:
                break
            
            # Bỏ qua nếu đoạn văn là heading 1,2,3
            if is_heading(para):
                exclusion_set.add(idx)
                continue
            
            prompt = f"""
                Viết lại đoạn văn sau bằng cách giảm từ '{word}'. 
                **TUYỆT ĐỐI GIỮ NGUYÊN MARKDOWN có sẵn: <<<PARA_START:Normal>>>, <<<PARA_START:List Bullet>>>**                

                Có thể cắt bỏ từ hoặc thay bằng từ đồng nghĩa hoặc thay đổi cấu trúc câu
                để giảm số lần xuất hiện của từ đó, miễn sao vẫn giữ nguyên ý nghĩa và văn phong tự nhiên.
                Không thêm hoặc bớt thông tin. 
                Chỉ trả về đoạn văn đã chỉnh sửa, gửi bản tốt nhất ko kèm lựa chọn.
                {para.text}"""
            
            new_text = call_chatbot(prompt, model_name, model_type, api_key)
            
            attempts[idx] = attempts.get(idx, 0) + 1
            if attempts[idx] >= 5:
                exclusion_set.add(idx)
            
            update_paragraph_text(para, new_text)
            current_density = get_overall_word_density_in_doc(word, doc)
            
    doc = split_sections(doc)
    return doc


def update_paragraph_text(paragraph, new_text):
    """
    Cập nhật nội dung của đoạn văn, giữ nguyên style và định dạng paragraph-level.
    Lưu ý: Các run cũ sẽ được giữ lại các định dạng (như in đậm, in nghiêng) và chỉ thay đổi nội dung.
    """
    # Xóa hết nội dung trong paragraph nhưng vẫn giữ lại các run cũ
    for run in paragraph.runs:
        run.text = ""  # Xóa nội dung trong các run cũ, giữ định dạng.

    # Thêm nội dung mới vào các run cũ
    # Đây là cách bạn thêm vào một run mới, nếu cần thiết có thể chia đoạn văn thành nhiều run
    paragraph.add_run(new_text)


def contains_long_sentence(text, min_words):
    """
    Kiểm tra xem text có chứa ít nhất 1 câu có số từ vượt quá min_words không.
    Tách câu theo dấu chấm, hỏi, cảm.
    """
    sentences = re.split(r'(?<=[.!?])\s+', text)
    for sentence in sentences:
        if len(sentence.split()) > min_words:
            return True
    return False

def process_long_sentences(doc, model_name, model_type, api_key, min_words, max_attempts=5):
    for para in doc.paragraphs:
        # Bỏ qua Heading hoặc đoạn in nghiêng
        if para.style.name.startswith("Heading") or any(run.italic for run in para.runs):
            continue



        attempts = 0
        while attempts < max_attempts and contains_long_sentence(para.text, min_words):
            prompt = f"""

            Viết lại đoạn văn sau để giảm độ dài của các câu vượt quá {min_words} từ,
            chỉ thay đổi cấu trúc và rút gọn các câu quá dài mà không làm mất ý nghĩa.
            Nếu câu dài có thể tách làm 2 câu ngắn, cho phép tái cấu trúc câu, viết dễ hiểu nhất có thể.
            Không thêm thông tin, không sử dụng markdown, không có lời dẫn, không có tiêu đề.
            
            {para.text}"""
            
            new_text = call_chatbot(prompt, model_name, model_type, api_key)

            update_paragraph_text(para, new_text)
            attempts += 1
    return doc


def key_brand_reduction(doc, brand_key, max_key, model_name, model_type, api_key):
    """
    Hàm này giảm số lần xuất hiện của brand_key trong tài liệu nếu số lần xuất hiện vượt quá max_key.
    """
    doc = merge_sections(doc)    

    import re
    import random


    # Hàm kiểm tra đoạn văn chỉ có một câu
    def is_excluded_paragraph(index):
        excluded_indices = set()

        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()  # Lấy văn bản của đoạn và loại bỏ khoảng trắng thừa
            if not para_text:
                continue  # Bỏ qua các đoạn văn trống

            # Đếm số dấu chấm trong đoạn văn
            sentence_count = len(re.findall(r'\.', para_text))

            # Kiểm tra nếu đoạn văn chỉ có 1 câu (1 dấu chấm hoặc không có dấu chấm)
            if sentence_count == 1 or sentence_count == 0:
                excluded_indices.add(i)

        return index in excluded_indices

    def find_paragraph_with_max_keyword(doc, brand_key, start_idx, end_idx):
        """
        Tìm đoạn văn có số lượng từ khóa xuất hiện nhiều nhất trong khoảng chỉ mục cho phép.
        Nếu có nhiều đoạn bằng nhau, ưu tiên chọn ngẫu nhiên thay vì theo thứ tự.
        """
        max_count = 0
        max_paragraphs = []  # List to store paragraphs with the max keyword count
        
        for i in range(start_idx, end_idx):
            if is_excluded_paragraph(i):  # Kiểm tra xem đoạn văn có cần bỏ qua không
                continue  # Nếu đoạn văn nằm trong danh sách bỏ qua thì tiếp tục vòng lặp mà không xử lý đoạn này

            paragraph = doc.paragraphs[i]
            count = paragraph.text.lower().count(brand_key.lower())
            
            if count > max_count:
                # Found a new max count, reset the list
                max_count = count
                max_paragraphs = [i]
            elif count == max_count:
                # If the count is equal, add this paragraph to the list
                max_paragraphs.append(i)

        # Select a random paragraph index from the max_paragraphs list
        if max_paragraphs:
            max_index = random.choice(max_paragraphs)
            return max_index, max_count
        else:
            return -1, 0  # Return a default value if no paragraphs found


    # Đếm tổng số lần xuất hiện của key_word trong toàn bộ tài liệu
    total_count = count_keywords_in_document(doc, brand_key)
    total_paragraphs = len(doc.paragraphs)
    start_idx = 0  # Vị trí bắt đầu xử lý
    end_idx = total_paragraphs  # Vị trí kết thúc (paragraph thứ 2 từ dưới lên)
    api_call_count = 0

    # Kiểm tra nếu số lần xuất hiện của key_word > max_key thì tiến hành xử lý
    while total_count > max_key:
        max_index, max_count = find_paragraph_with_max_keyword(doc, brand_key, start_idx, end_idx)

        if max_index == -1 or api_call_count>=50:  # Không có đoạn văn phù hợp để xử lý
            break

        paragraph = doc.paragraphs[max_index]

        prompt = f"""
        **Từ khóa:** {brand_key}
        **TUYỆT ĐỐI GIỮ NGUYÊN MARKDOWN có sẵn: <<<PARA_START:Normal>>>, <<<PARA_START:List Bullet>>>**                
        **Đây là từ khóa thương hiệu, có thể thay thành 'nền tảng', 'chúng tôi', 'thương hiệu', 'nhà cái' ...**
        **Đoạn văn cần xử lý:** {paragraph.text}
        **Hướng dẫn:**

        *   **Giảm từ khóa:** Giảm từ khóa {brand_key} xuống còn 1. 
        *   **Tính logic:** Đảm bảo đoạn văn sau khi chỉnh sửa vẫn có tính logic và mạch lạc.
        *   **Dễ hiểu:** Sử dụng ngôn ngữ dễ hiểu, tự nhiên sẵn có.
        *   **Không định dạng:** Chỉ trả về đoạn văn đã chỉnh sửa, không kèm theo lời dẫn, markdown.
        """


        new_text = call_chatbot(prompt, model_name, model_type, api_key)
        
        api_call_count += 1


        # Kiểm tra và xử lý nếu new_text không hợp lệ
        if not isinstance(new_text, str) or not new_text.strip():
            print(f"Lỗi: Kết quả trả về từ call_chatgpt không hợp lệ. Giữ nguyên đoạn văn bản gốc.")
            new_text = paragraph.text  # Giữ nguyên đoạn văn bản gốc nếu không hợp lệ


        # Lưu lại tất cả các runs gốc trong paragraph
        runs = paragraph.runs

        # Xóa toàn bộ nội dung cũ trong paragraph
        paragraph.clear()

        # Thêm đoạn văn bản mới
        run = paragraph.add_run(new_text)

        # Sao chép lại tất cả các thuộc tính định dạng từ các run cũ vào run mới
        for orig_run in runs:
            run.font.bold = orig_run.bold
            run.font.italic = orig_run.italic
            run.font.underline = orig_run.underline
            run.font.size = orig_run.font.size
            run.font.name = orig_run.font.name
            run.font.color.rgb = orig_run.font.color.rgb

        # Cập nhật lại tổng số từ khóa trong tài liệu
        total_count = count_keywords_in_document(doc, brand_key)
        print(f"Kết thúc lần xử lý, số lần gọi API: {api_call_count}, tổng số từ khóa còn lại: {total_count}")
        

    doc = split_sections(doc)

    return doc


def remove_double_spaces(doc):
    """
    Hàm xử lý để loại bỏ tất cả double space (hoặc nhiều khoảng trắng liên tiếp)
    trong nội dung của Document, giữ nguyên định dạng.
    
    :param doc: Đối tượng Document đã được đọc bởi python-docx
    :return: Đối tượng Document sau khi đã xóa double spaces
    """
    for paragraph in doc.paragraphs:
        # 1) Xử lý bên trong từng run, thay tất cả chuỗi "  " (>=2 dấu cách) thành 1 dấu cách
        for run in paragraph.runs:
            while "  " in run.text:
                run.text = run.text.replace("  ", " ")
        
        # 2) Sau khi đã thay thế double space trong mỗi run,
        #    ta kiểm tra khoảng trống đầu/cuối run để tránh double space giữa các run
        i = 0
        while i < len(paragraph.runs) - 1:
            current_run = paragraph.runs[i]
            next_run = paragraph.runs[i + 1]
            
            # Nếu run hiện tại kết thúc bằng space và run kế tiếp bắt đầu bằng space => xóa 1 space ở đầu run kế tiếp
            if current_run.text.endswith(" ") and next_run.text.startswith(" "):
                next_run.text = next_run.text[1:]  # Bỏ bớt 1 space ở đầu
            else:
                i += 1
    
    return doc


import re
from docx import Document

def copy_run_style(source_font, target_font):
    """
    Sao chép một số thuộc tính cơ bản của font từ run gốc sang run mới.
    Ở đây chỉ copy những thuộc tính thường dùng, nếu cần thêm thuộc tính,
    bạn có thể bổ sung tương tự.
    """
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.color.rgb = source_font.color.rgb
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline

def is_bullet_paragraph(paragraph):
    """
    Ví dụ kiểm tra paragraph có phải bullet list hay không bằng cách
    đơn giản xem tên style có chứa 'List' hay không.
    Thực tế bạn có thể tùy chỉnh logic này cho phù hợp.
    """
    if paragraph.style and paragraph.style.name:
        return "List" in paragraph.style.name
    return False

def bold_text_before_colon(paragraph, colon_index):
    """
    In đậm toàn bộ đoạn text (theo các runs) nằm trước dấu ':' đầu tiên,
    giữ nguyên định dạng gốc cho phần sau dấu ':' cũng như các định dạng
    (màu, font, v.v.) trước đó. Chỉ thêm bold cho phần trước dấu ':'.
    """
    chars_processed = 0
    for run in paragraph.runs:
        run_text = run.text
        run_len = len(run_text)

        # Nếu đã xử lý xong (vượt quá vị trí dấu ':') thì thoát
        if chars_processed >= colon_index:
            break

        # Trường hợp cả run nằm trước dấu ':'
        if chars_processed + run_len <= colon_index:
            run.font.bold = True  # chỉ thêm bold, giữ nguyên màu, font,...
        else:
            # Dấu ':' nằm giữa run này => cần "chẻ" run
            before_length = colon_index - chars_processed
            before_text = run_text[:before_length]
            colon_char = run_text[before_length]  # chính là ':'
            after_text = run_text[before_length+1:]  # phần còn lại sau ':'

            old_font = run.font
            # Xóa nội dung gốc của run hiện tại (để ta tạo lại 3 run)
            run.text = ""

            # 1) Run trước dấu ':'
            if before_text:
                run_before = paragraph.add_run(before_text)
                copy_run_style(old_font, run_before.font)
                run_before.font.bold = True  # chỉ thêm bold cho phần trước dấu ':'

            # 2) Run cho phần dấu ':'
            run_colon = paragraph.add_run(colon_char)
            copy_run_style(old_font, run_colon.font)
            # Dấu ':' không in đậm (tùy bạn, nếu muốn in đậm luôn thì set = True)

            # 3) Run cho phần sau dấu ':'
            if after_text:
                run_after = paragraph.add_run(after_text)
                copy_run_style(old_font, run_after.font)
                # Không can thiệp bold ở đây, giữ nguyên font.bold ban đầu

            # Sau khi xử lý xong vị trí dấu ':', không cần xử lý nữa
            break

        chars_processed += run_len

def bold_bullet_list(doc):
    """
    Duyệt các paragraph trong Document:
      - Nếu là bullet list và có dấu ':' ở cuối (không còn nội dung nào sau ':'),
        chuyển style sang Normal (thành văn bản thường).
      - Nếu là bullet list và có dấu ':' ở giữa, in đậm phần trước dấu ':',
        giữ nguyên các định dạng khác.
    Trả về doc sau khi đã xử lý.
    """
    for paragraph in doc.paragraphs:
        if is_bullet_paragraph(paragraph):
            text = paragraph.text
            colon_index = text.find(':')

            if colon_index == -1:
                # Không có dấu ':' => không làm gì
                continue

            # Nếu ':' là ký tự cuối cùng (và sau đó không có nội dung gì thêm)
            if colon_index == len(text) - 1:
                # Chuyển style sang Normal
                paragraph.style = doc.styles['Normal']
            else:
                # In đậm phần trước dấu ':'
                bold_text_before_colon(paragraph, colon_index)

    return doc


def process_file_normalize(file_path, output_folder, max_key, number_of_image, model_name, model_type, api_key, language, word_density, brand_key = None):
    try:  
        doc = Document(file_path)
        file_name = os.path.basename(file_path)
        key_word = file_name.replace('.docx', '').strip().lower()
        print(f"Processing file: {file_path}")
        
        if brand_key:
            print('Đang giảm brand key')
            doc = key_brand_reduction(doc, brand_key, max_key - 4, model_name, model_type, api_key)

        print('đang xử lý mật độ')
        doc = process_word_density_doc(doc, key_word, model_name, model_type, api_key, word_density) #api
 
        print('đang xử lý câu dài')
        doc = process_long_sentences(doc, model_name, model_type, api_key, min_words=30) #api

        print('xử lý xong mật độ và câu dài')

        # doc = fix_meta(doc, key_word, model_name, model_type, api_key, language)  #api


        doc = check_intro(doc, key_word, model_name, model_type, api_key, language)   #api
  
        doc = check_conclusion(doc, key_word, model_name, model_type, api_key, language)   #api


        print('xử lý xong intro/conclusion')


        if count_keywords_in_document(doc, key_word) > max_key:
            doc = key_reduction(doc, key_word , max_key - 5, model_name, model_type, api_key, language)   #api
        else:

            doc = insert_keyword_below_h3(doc, key_word, max_key - 5, model_name, model_type, api_key, language)   #api
            
            doc = insert_keyword(doc, key_word, max_key - 5, model_name, model_type, api_key, language)   #api
            

        print('xử lý xong keyword')


        doc = generate_h1(doc, key_word, model_name, model_type, api_key, language)  #api
 
        doc = generate_h3(doc, key_word, model_name, model_type, api_key, language)  #api
        
        doc = generate_h2(doc, key_word, model_name, model_type, api_key, language)   #api


        doc = copy_and_insert_h3(doc, key_word, number_of_image, model_name, model_type, api_key, language)  #insert tên ảnh auto

        doc = normalize_headings(doc, model_name, model_type, api_key, language)  #api lower
         

        doc = remove_double_spaces(doc)

        # Đảm bảo thư mục output tồn tại
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # Lưu file đã được cập nhật vào thư mục con
        output_file_path = os.path.join(output_folder, f"{os.path.basename(file_path)}")
        
        doc.save(output_file_path)
        
        doc = Document(output_file_path)
                
        doc = justify_and_set_line_spacing(doc)

        doc = format_image_names_in_document(doc)


        doc = change_font_to_arial(doc)
        
        doc = bold_bullet_list(doc)
        # Lưu tài liệu mới
        doc.save(output_file_path)
        

        print(f"Đã lưu kết quả vào file: {output_file_path}")
    
    except Exception as e:
        print(f"Lỗi xử lý file {file_path}: {e}")
