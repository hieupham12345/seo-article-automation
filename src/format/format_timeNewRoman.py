#format time new roman, ...
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

def justify_and_set_line_spacing(doc):
    try:
        for i, paragraph in enumerate(doc.paragraphs):
            if i != 1:  # Đảm bảo rằng đoạn văn số 2 (index 1) không bị thay đổi
                # Căn đều đoạn văn
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Cài đặt giãn cách dòng 1.5
                paragraph.paragraph_format.line_spacing = 1.5
                
                # Cài đặt giãn cách đoạn trước và sau
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(10)
        
        return doc

    except Exception as e:
        return f"Lỗi xử lý căn đều và giãn cách dòng: {e}"
    

def normalize_image(doc):
    try:
        num_paragraphs = len(doc.paragraphs)


        for i in range(num_paragraphs):
            paragraph = doc.paragraphs[i]
            style_name = paragraph.style.name



            if style_name == 'Normal':
                has_italic = any(run.font.italic for run in paragraph.runs)

                if has_italic:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Căn giữa nếu có in nghiêng

        return doc
    except Exception as e:
        return f"Lỗi xử lý file: {e}"


def format_run(run, font_name="Times New Roman", font_size=None):
    """Định dạng một đoạn văn bản với font và giữ nguyên các thuộc tính cũ."""
    original_bold = run.font.bold
    original_italic = run.font.italic
    original_underline = run.font.underline
    original_color = run.font.color.rgb
    
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if original_bold:
        run.font.bold = True
    if original_italic:
        run.font.italic = True
    if original_underline:
        run.font.underline = original_underline
    if original_color:
        run.font.color.rgb = original_color

    run.font.color.rgb = RGBColor(0, 0, 0)

def process_doc_time_new_roman(doc):
    """Chỉnh font toàn bộ văn bản và chuẩn hóa heading."""
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1":
            for run in paragraph.runs:
                format_run(run, font_size=20)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif paragraph.style.name == "Heading 2":
            for run in paragraph.runs:
                format_run(run, font_size=16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif paragraph.style.name == "Heading 3":
            for run in paragraph.runs:
                format_run(run, font_size=14)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            for run in paragraph.runs:
                format_run(run, font_size=14)


                    
        # Kiểm tra nếu đoạn tiếp theo là đoạn văn bình thường ngay sau Heading 2
        if paragraph.style.name == "Heading 2" and i + 1 < len(doc.paragraphs):
            next_paragraph = doc.paragraphs[i + 1]
            if next_paragraph.style.name not in ["Heading 1", "Heading 2", "Heading 3"]:
                for run in next_paragraph.runs:
                    format_run(run, font_size=14)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        format_run(run, font_size=14)
    
    return doc


def process_folder_format_time(input_folder):
    """Xử lý hàng loạt file trong thư mục input_folder và ghi đè file cũ."""
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            file_path = os.path.join(input_folder, filename)
            

            doc = Document(file_path)
            
            doc = process_doc_time_new_roman(doc)
            
            doc = justify_and_set_line_spacing(doc)

            doc = normalize_image(doc)
            doc.save(file_path)
            print(f"Processed: {filename}")