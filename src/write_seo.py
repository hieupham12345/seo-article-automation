from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os
import re
import traceback
import json
from src.chat_bot import call_chatbot

# Hàm insert_paragraph_after (improved)
def insert_paragraph_after(paragraph, text="", style=None, font_size=None, bold=False, italic=False, alignment=None):
    try:
        if not text.strip():  
            raise ValueError("Nội dung đoạn văn bị rỗng, không thể chèn.")  

        new_p = OxmlElement('w:p')
        if new_p is None:
            raise RuntimeError("Không thể tạo phần tử đoạn văn mới.")  

        paragraph._element.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)

        if not new_paragraph:
            raise RuntimeError("Không thể thêm đoạn văn mới vào tài liệu.")  

        if style:
            new_paragraph.style = style

        run = new_paragraph.add_run(text.strip())
        if font_size:
            run.font.size = font_size
        run.bold = bold
        run.italic = italic
        if alignment:
            new_paragraph.alignment = alignment

        return new_paragraph  

    except Exception as e:
        raise RuntimeError(f"Lỗi khi chèn đoạn văn: {e}")  

def clean_empty_paragraphs(doc):
    """Removes truly empty paragraphs (after removing whitespace and non-text content) without affecting formatting."""
    paragraphs_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text_content = re.sub(r'\s+', '', p.text)  # remove all whitespace
        if not text_content:  # Check if paragraph is actually empty after removing whitespace
            paragraphs_to_remove.append(i)

    # Remove paragraphs in reverse order to avoid index issues during removal
    for i in reversed(paragraphs_to_remove):
        try:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        except IndexError:
            print(f"Warning: Could not remove paragraph at index {i}. This might be due to previous paragraph removal.")

def is_heading_3(para):
    """Kiểm tra có phải là Heading 3 không"""
    return para.style.name == 'Heading 3'


# Hàm kiểm tra tiêu đề H3
def is_heading_2(para):
    return para.style.name == 'Heading 2'

# Hàm kiểm tra nếu file đã tồn tại trong thư mục đích
def is_file_processed(file_path, output_folder):
    output_file_path = os.path.join(output_folder, os.path.basename(file_path))
    return os.path.exists(output_file_path)


def clean_json(data):
    try:
        # Chuyển thành JSON string rồi parse lại để kiểm tra
        json_str = json.dumps(data, ensure_ascii=False)
        cleaned_data = json.loads(json_str)
        return cleaned_data
    except json.JSONDecodeError as e:
        print("Lỗi JSON:", e)
        return None
    
def convert_to_dict(raw_data):
    """
    Robust function to convert a string to a dictionary, handling various markdown formats.
    
    Args:
        raw_data (str): The input string potentially containing markdown or extra formatting
    
    Returns:
        dict or None: Parsed dictionary or None if parsing fails
    """
    # Remove markdown code block markers if present
    cleaned_data = raw_data.strip()
    
    # Remove ```json or ```python or other code block markers
    if cleaned_data.startswith('```'):
        cleaned_data = cleaned_data.split('\n', 1)[-1]
    
    # Remove trailing code block marker
    if cleaned_data.endswith('```'):
        cleaned_data = cleaned_data[:-3]
    
    # Remove any leading/trailing whitespace
    cleaned_data = clean_json(cleaned_data).strip()

    
    try:
        # Attempt to parse the cleaned data as JSON
        data_dict = json.loads(cleaned_data)
        return data_dict
    except json.JSONDecodeError as e:
        print(f"Error converting string to dictionary: {e}")
        print(f"Problematic input: {cleaned_data}")
        return False

def fix_text_between_h2_h3(doc, model_name, model_type, api_key, key_word, language):
    """
    Xử lý tài liệu Word để chèn văn bản giữa H2 và H3 nếu sau H2 là H3.
    """
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if para.style.name == 'Heading 2':
            h2_para = para
            if i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]
                if next_para.style.name == 'Heading 3':
                    h2_text = h2_para.text.strip()
                    

                    prompt = f"""
                        Viết 2 câu giới thiệu mở đầu cho đoạn nội dung thuộc heading: "{h2_text}".  
                        Nội dung phải viết **hoàn toàn bằng {language}**, không chứa bất kỳ ngôn ngữ khác.  
                        Không lặp lại "{h2_text}" trong câu, không chèn "{key_word}".  
                        Không sử dụng từ "heading" hoặc đề cập đến "bài viết này".  
                        Chỉ trả về 2 câu trong cùng một đoạn, không có tiêu đề hay markdown.  
                        Tối đa 250 ký tự.  

                    """


                    
                    inserted_text = call_chatbot(prompt, model_name, model_type, api_key)

                    if inserted_text:
                        inserted_text = inserted_text.strip()
                        if len(inserted_text) > 0:
                            inserted_text = inserted_text[0].upper() + inserted_text[1:]

                        new_para = insert_paragraph_after(
                            paragraph=h2_para,
                            text=inserted_text,
                            style='Normal',
                            font_size=Pt(11),
                            bold=False,
                            italic=False,
                            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        )
                        i += 1
        i += 1
    return doc



def convert_to_list_format(doc):
    # Duyệt qua các paragraph trong tài liệu
    for para in doc.paragraphs:
        # Kiểm tra nếu paragraph bắt đầu bằng dấu "*"
        if para.text.strip().startswith('*'):
            # Loại bỏ dấu '*' trước khi chuyển thành danh sách
            para.text = para.text.lstrip('*').lstrip()  # Xóa dấu * và khoảng trắng sau dấu *
            
            # Thay đổi style của paragraph thành danh sách
            para.style = 'List Bullet'  # Sử dụng style danh sách có gạch đầu dòng
    return doc

def extract_word_to_json(doc):
    """
    Extracts data from a Word document and formats it as a JSON object,
    stopping when encountering the last Heading 2 with text "Kết luận".
    
    Args:
        doc: A Document object from the python-docx library.
    
    Returns:
        dict: A JSON-formatted dictionary with heading 2, heading 3, and their corresponding content.
    """
    result = []
    current_h2 = None
    current_h3 = None

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 2'):
            # Stop processing if the Heading 2 is "Kết luận"
            if paragraph.text.strip() == "Kết luận":
                break
            
            # Save the current heading 2 and reset heading 3
            current_h2 = {
                "heading_2": paragraph.text.strip(),
                "content": [],
            }
            result.append(current_h2)

        elif paragraph.style.name.startswith('Heading 3') and current_h2:
            # Save the current heading 3 under the current heading 2
            current_h3 = {
                "heading_3": paragraph.text.strip(),
                "content": []
            }
            current_h2["content"].append(current_h3)

        elif paragraph.style.name == 'Normal' and paragraph.text.strip():
            # Add normal paragraph to the appropriate level
            if current_h3:  # Add to the last heading 3
                current_h3["content"].append(paragraph.text.strip())
            elif current_h2:  # Add directly to the heading 2
                current_h2["content"].append(paragraph.text.strip())

    return result


def insert_paragraph_after_for_split_para(paragraph: Paragraph, text: str = None, style: str = None) -> Paragraph:
    """
    Chèn một paragraph mới ngay sau paragraph hiện tại.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def delete_paragraph(paragraph: Paragraph):
    """
    Xóa paragraph khỏi tài liệu.
    """
    p_element = paragraph._element
    p_element.getparent().remove(p_element)

def split_sentences(text: str):
    """
    Tách văn bản thành danh sách các câu.
    Dùng regex tách theo dấu chấm, dấu chấm than, dấu hỏi và khoảng trắng sau đó.
    """
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s for s in sentences if s]

def group_sentences(sentences, max_words=80):
    """
    Gom các câu thành các đoạn sao cho mỗi đoạn chứa tối đa max_words từ.
    Ưu tiên gom từ trên xuống: đoạn trước càng đầy (tới 80 từ) thì đoạn sau nhận phần còn lại.
    """
    grouped_paras = []
    current_sentences = []
    current_count = 0

    for sentence in sentences:
        word_count = len(sentence.split())
        if current_count + word_count > max_words and current_sentences:
            grouped_paras.append(" ".join(current_sentences))
            current_sentences = [sentence]
            current_count = word_count
        else:
            current_sentences.append(sentence)
            current_count += word_count
    if current_sentences:
        grouped_paras.append(" ".join(current_sentences))
    return grouped_paras

def split_paragraho_doc(doc):
    """
    Cách tiếp cận 2-phase:
      Phase 1: Quét qua doc, tìm Heading 3 và các Normal đi kèm, lưu kết quả.
      Phase 2: Xóa các đoạn cũ, chèn các đoạn mới sau Heading 3 tương ứng.
    """
    paras = list(doc.paragraphs)

    # Lưu trữ thông tin để xử lý về sau.
    # Mỗi phần tử sẽ có dạng:
    # (heading_3_paragraph, [list_of_normal_paragraphs_to_combine])
    heading_info = []
    
    i = 0
    while i < len(paras):
        p = paras[i]
        if p.style.name == "Heading 3":
            # Bắt đầu gom các paragraph Normal tiếp theo
            group_indices = []
            j = i + 1
            while j < len(paras):
                p_next = paras[j]
                if p_next.style.name in ["Heading 2", "Heading 3", "List Bullet"]:
                    break
                if p_next.style.name == "Normal":
                    group_indices.append(paras[j])
                j += 1
            
            # Nếu có gom được đoạn normal, ta lưu lại thông tin
            if group_indices:
                heading_info.append((p, group_indices))
            i = j
        else:
            i += 1
    
    # PHASE 2: Thực hiện xóa và chèn
    # Lưu ý: nên làm từ cuối danh sách heading_info ngược lên để không ảnh hưởng
    # đến vị trí paragraph trong doc khi xóa/chèn.
    for heading_para, normal_paras in reversed(heading_info):
        # Kết hợp tất cả các paragraph Normal
        combined_text = " ".join(np.text.strip() for np in normal_paras if np.text.strip())
        # Tách thành câu, rồi group thành các đoạn < 80 từ
        sentences = split_sentences(combined_text)
        new_text_paras = group_sentences(sentences, max_words=80)
        
        # Xóa các paragraph Normal cũ
        for np in normal_paras:
            delete_paragraph(np)
        
        # Chèn các đoạn mới ngay sau Heading 3
        last_para = heading_para
        for new_para_text in new_text_paras:
            inserted_para = insert_paragraph_after_for_split_para(last_para, text=new_para_text, style="Normal")
            last_para = inserted_para  # Cập nhật điểm "neo" để chèn lần kế
    
    return doc



def remove_special_chars_from_docx(doc):
    
    for para in doc.paragraphs:
        para.text = para.text.replace("'", "").replace('"', "").replace("*", "").replace("“", "").replace("”", "")
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("'", "").replace('"', "").replace("*", "")
    
    return doc

def format_run(run, font_name="Arial", font_size=None):
    """Định dạng một đoạn văn bản với font và giữ nguyên các thuộc tính cũ."""
    original_bold = run.font.bold
    original_italic = run.font.italic
    original_underline = run.font.underline
    original_color = run.font.color.rgb
    
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if original_bold:
        run.font.bold = True
    if original_italic:
        run.font.italic = True
    if original_underline:
        run.font.underline = original_underline
    if original_color:
        run.font.color.rgb = original_color

    run.font.color.rgb = RGBColor(0, 0, 0)


def process_doc_arial(doc):
    """Chỉnh font toàn bộ văn bản và chuẩn hóa heading."""
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1":
            for run in paragraph.runs:
                format_run(run, font_size=20)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif paragraph.style.name == "Heading 2":
            for run in paragraph.runs:
                format_run(run, font_size=17)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif paragraph.style.name == "Heading 3":
            for run in paragraph.runs:
                format_run(run, font_size=13)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            for run in paragraph.runs:
                format_run(run, font_size=11)


                    
        # Kiểm tra nếu đoạn tiếp theo là đoạn văn bình thường ngay sau Heading 2
        if paragraph.style.name == "Heading 2" and i + 1 < len(doc.paragraphs):
            next_paragraph = doc.paragraphs[i + 1]
            if next_paragraph.style.name not in ["Heading 1", "Heading 2", "Heading 3"]:
                for run in next_paragraph.runs:
                    format_run(run, font_size=11)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        format_run(run, font_size=11)
    
    return doc

def justify_and_set_line_spacing(doc):
    try:
        for i, paragraph in enumerate(doc.paragraphs):
            if i != 1:  # Đảm bảo rằng đoạn văn số 2 (index 1) không bị thay đổi
                # Căn đều đoạn văn
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Cài đặt giãn cách dòng 1.5
                paragraph.paragraph_format.line_spacing = 1.5
                
                # Cài đặt giãn cách đoạn trước và sau
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(10)
        
        return doc

    except Exception as e:
        return f"Lỗi xử lý căn đều và giãn cách dòng: {e}"


import re
from docx import Document

def copy_run_style(source_font, target_font):
    """
    Sao chép một số thuộc tính cơ bản của font từ run gốc sang run mới.
    Ở đây chỉ copy những thuộc tính thường dùng, nếu cần thêm thuộc tính,
    bạn có thể bổ sung tương tự.
    """
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.color.rgb = source_font.color.rgb
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline

def is_bullet_paragraph(paragraph):
    """
    Ví dụ kiểm tra paragraph có phải bullet list hay không bằng cách
    đơn giản xem tên style có chứa 'List' hay không.
    Thực tế bạn có thể tùy chỉnh logic này cho phù hợp.
    """
    if paragraph.style and paragraph.style.name:
        return "List" in paragraph.style.name
    return False

def bold_text_before_colon(paragraph, colon_index):
    """
    In đậm toàn bộ đoạn text (theo các runs) nằm trước dấu ':' đầu tiên,
    giữ nguyên định dạng gốc cho phần sau dấu ':' cũng như các định dạng
    (màu, font, v.v.) trước đó. Chỉ thêm bold cho phần trước dấu ':'.
    """
    chars_processed = 0
    for run in paragraph.runs:
        run_text = run.text
        run_len = len(run_text)

        # Nếu đã xử lý xong (vượt quá vị trí dấu ':') thì thoát
        if chars_processed >= colon_index:
            break

        # Trường hợp cả run nằm trước dấu ':'
        if chars_processed + run_len <= colon_index:
            run.font.bold = True  # chỉ thêm bold, giữ nguyên màu, font,...
        else:
            # Dấu ':' nằm giữa run này => cần "chẻ" run
            before_length = colon_index - chars_processed
            before_text = run_text[:before_length]
            colon_char = run_text[before_length]  # chính là ':'
            after_text = run_text[before_length+1:]  # phần còn lại sau ':'

            old_font = run.font
            # Xóa nội dung gốc của run hiện tại (để ta tạo lại 3 run)
            run.text = ""

            # 1) Run trước dấu ':'
            if before_text:
                run_before = paragraph.add_run(before_text)
                copy_run_style(old_font, run_before.font)
                run_before.font.bold = True  # chỉ thêm bold cho phần trước dấu ':'

            # 2) Run cho phần dấu ':'
            run_colon = paragraph.add_run(colon_char)
            copy_run_style(old_font, run_colon.font)
            # Dấu ':' không in đậm (tùy bạn, nếu muốn in đậm luôn thì set = True)

            # 3) Run cho phần sau dấu ':'
            if after_text:
                run_after = paragraph.add_run(after_text)
                copy_run_style(old_font, run_after.font)
                # Không can thiệp bold ở đây, giữ nguyên font.bold ban đầu

            # Sau khi xử lý xong vị trí dấu ':', không cần xử lý nữa
            break

        chars_processed += run_len

def bold_bullet_list(doc):
    """
    Duyệt các paragraph trong Document:
      - Nếu là bullet list và có dấu ':' ở cuối (không còn nội dung nào sau ':'),
        chuyển style sang Normal (thành văn bản thường).
      - Nếu là bullet list và có dấu ':' ở giữa, in đậm phần trước dấu ':',
        giữ nguyên các định dạng khác.
    Trả về doc sau khi đã xử lý.
    """
    for paragraph in doc.paragraphs:
        if is_bullet_paragraph(paragraph):
            text = paragraph.text
            colon_index = text.find(':')

            if colon_index == -1:
                # Không có dấu ':' => không làm gì
                continue

            # Nếu ':' là ký tự cuối cùng (và sau đó không có nội dung gì thêm)
            if colon_index == len(text) - 1:
                # Chuyển style sang Normal
                paragraph.style = doc.styles['Normal']
            else:
                # In đậm phần trước dấu ':'
                bold_text_before_colon(paragraph, colon_index)

    return doc

def process_word_file(file_path, main_key, secondary_keys, model_name, model_type, output_folder, api_key, language, seo_category, learn_data=None):
    try:

        doc = Document(file_path)

        output_file_path = os.path.join(output_folder, os.path.basename(file_path))

        h2_h3_map = {}
        current_h2 = None

        # Xây dựng bản đồ H2 và các H3 tương ứng
        for para in doc.paragraphs:
            if is_heading_2(para):
                current_h2 = para.text.strip()
                h2_h3_map[current_h2] = []
            elif is_heading_3(para) and current_h2:
                h2_h3_map[current_h2].append(para.text.strip())

        # Xử lý H2 cuối cùng nếu cần loại bỏ
        if current_h2 and not h2_h3_map[current_h2]:
            del h2_h3_map[current_h2]

        if learn_data:
            print(f'có learn_data: {learn_data[:10]}')
        else:
        # In ra để kiểm tra
            print(f"main_key: {main_key}")
                    
        # Gửi từng H2 cùng các H3 tương ứng đến chatbot và nhận lại JSON
        for h2, h3_list in h2_h3_map.items():
            data = {
                "H2": h2,
                "H3": h3_list,
            } 
            
            try:
                #tạm thời ko dùng
                data_seo = extract_word_to_json(doc)
                
                prompt_for_all = f"""
                    **DỮ LIỆU ĐẦU VÀO:**
                                                                    
                            * **NGÔN NGỮ VIẾT:** {language}.
                            * **TUYỆT ĐỐI KHÔNG ĐƯỢC THAY ĐỔI NGÔN NGỮ:** Toàn bộ nội dung PHẢI được viết HOÀN TOÀN bằng {language}.
                            * **KIỂM TRA NGÔN NGỮ:** Sau khi hoàn thành, BUỘC PHẢI kiểm tra lại 100% nội dung để đảm bảo chính xác.

                            * **Dữ Liệu Học Tập (ƯU TIÊN CAO NHẤT):** {learn_data} (Mở rộng và làm mới 100%, viết lại hoàn toàn tránh đạo văn, bổ sung chi tiết theo 5W1H, viết sáng tạo)
                            * **Dữ Liệu Cần Viết (CẤU TRÚC BÀI VIẾT):** {data} (Xây dựng cấu trúc H2, H3. Đây là nội dung sẽ viết)
                            * **Outline:** {h2_h3_map} (Tham khảo kỹ các heading, tránh lặp nội dung hay lan man sang heading khác)
                            
                            * **Từ Khóa Chính:** {main_key}
                            * **Từ Khóa Phụ:** {secondary_keys} (bổ nghĩa cho bài) 

                            **YÊU CẦU NỘI DUNG (CHUẨN EAT):**

                            * **Phạm vi:** CHỈ viết nội dung H3 hiện tại, KHÔNG đề cập H3 khác.
                            * **EEAT:** Chuyên môn cao, chính xác, chi tiết, đầy đủ, dễ hiểu. Luôn phải viết cụ thể và đầy đủ, không bỏ lưng chừng.
                                * Khi giải thích về quy luật, luật chơi, khái niệm, cách thức hướng dẫn cần viết đầy đủ từ A-Z.  
                                * Khi giới thiệu doanh nghiệp/tổ chức/cá nhân viết chuẩn EEAT (được phép tự tạo thông tin theo chuẩn mực ngành).
                                * **Thông tin bạn tạo ra được coi là chính xác hoàn toàn, hãy dẫn dắt tự nhiên**
                            * **5W1H:** Đầy đủ, dẫn chứng hợp lý, đảm bảo 100% có dẫn chứng.
                            * **Tránh chung chung:** Rõ ràng, dễ hiểu, thuyết phục. Không đặt câu hỏi.
                            * **Thái độ:** Tích cực, khẳng định thương hiệu, thể hiện niềm tự hào và sự tin tưởng vào giá trị sản phẩm/dịch vụ.
                            * **Tránh lan man:** Tập trung từ khóa chính.
                            * **Không lặp nội dung (Quan trọng nhất):** Dựa vào dữ liệu bài viết hiện tại, tuyệt đối không lặp các nội dung đã có ở heading khác.
                            
                            **DANH SÁCH:**
                            **SỬ DỤNG KHI CẦN LIỆT KÊ, QUY TRÌNH, HƯỚNG DẪN TỪNG BƯỚC. NỘI DUNG CÁC ĐOẠN SEO VÀ LIST KHÔNG ĐƯỢC LẶP LẠI GIỐNG NHAU**
                            **VIẾT NGẮN GỌN**
                            * Sử dụng dấu hoa thị (*) trong markdown ở đầu mỗi mục của list.
                            * Mỗi mục liệt kê cần viết ngắn gọn, đầy đủ thông tin theo nguyên tắc 5W1H.
                            * Mỗi mục chỉ nên trình bày trên một dòng để đảm bảo rõ ràng, dễ đọc và tạo sự thay đổi về độ dài dòng trong văn bản (Burstiness).
                            * Bắt buộc phải có tiêu đề list trước khi bắt đầu danh sách để tạo sự đa dạng cấu trúc câu (Burstiness).
                            * Tiêu đề list **KHÔNG** được sử dụng MARKDOWN '*' và không chứa '{main_key}', không dược viết hoa tiêu đề list.
                            
                            **TIÊU CHUẨN SEO:**
                            * **Độc đáo:** Tránh đạo văn, copy.
                            * **KHÔNG markdown in đậm trong bài**.
                            * **Mật độ từ khóa chính:** 1 LẦN/HEADING 3.
                            * **KHÔNG dùng câu hỏi.**
                            * **Độ dài mỗi heading 3**: **TỐI ĐA 100 CHỮ (KHÔNG TÍNH DANH SÁCH)**.
                            
                            **ĐỊNH DẠNG ĐẦU RA (JSON):**
                            * **Chỉ 1 kết quả JSON.**
                            * **KO GỬI VỀ DẠNG LIST**
                            * **Nếu có dấu " trong nội dung, hãy escape nó thành \\" để đảm bảo JSON hợp lệ.**
                            * **Key:** Các tiêu đề H3.
                            * **Value:** Nội dung (chuỗi - string).
                            * **NGẮT ĐOẠN:** CHỈ dùng markdown \\n, không xuống dòng bằng 'enter' tránh bị lỗi.
                            * **Không thêm lời dẫn.**
                            * **Viết đủ H3.**
                            * **KHÔNG dùng ký tự [ và ] trong nội dung.
                            Ví dụ json:
                            {{
                            "H3 title 1": "Paragrahp 1\\n Paragrahp 2\\n List\\n[Anchor text](url)",
                            "H3 title 2": "Paragrahp 1\\n Paragrahp 2\\n List",
                            ...
                            }}
                """
            
                if seo_category == "Storytelling":
                    prompt = f"""
                        * **QUY TẮC VĂN PHONG**

                            **Bắt đầu mộc mạc, gần gũi**: Hãy như kể chuyện phiếm, đừng trịnh trọng. Thêm chút tiếng lóng hoặc ví von tinh tế để tạo không khí.
                            **Nhịp điệu đa dạng**: Thay đổi tốc độ, khi thì dốc ngược bất ngờ, lúc lại chậm rãi để giữ hứng thú cho người đọc.
                            **Cảm xúc đan xen**: Nhẹ nhàng khi cần lặng trôi, sôi nổi lúc cao trào, tránh đều đều.  
                            **Chi tiết chân thực**: Kể cả những mẩu chuyện nhỏ, hương thơm vụt thoảng hay âm thanh rì rào, để khung cảnh thêm sống động.
                            **Giữ giọng thân thiện**: Đừng áp đặt suy nghĩ; gợi mở để độc giả tự cảm nhận và tham gia vào câu chuyện.
                            **Pha góc nhìn cá nhân**: Vừa nói lên quan điểm vừa để ngỏ chỗ cho người đọc “nửa tin nửa ngờ” và suy ngẫm.
                            **Tránh lặp từ**: Thay đổi cách diễn đạt và dùng từ đồng nghĩa. Hãy “chơi đùa” với ngôn ngữ để tạo màu sắc mới.
                            **Cụ thể hóa**: Vẽ bối cảnh tỉ mỉ, tránh lý thuyết mơ hồ. Đưa người đọc chạm đến hình ảnh rõ ràng.
                            **Ẩn dụ sáng tạo**: Tạo so sánh bất ngờ, thú vị, giúp câu văn thêm sinh động.
                            **Thêm màu văn hóa hoặc xu hướng**: Tận dụng ngôn ngữ vùng miền hay “hot trend” nếu hợp cảnh, để bài viết tươi mới.
                            **Pha trộn nghiêm túc và hài hước**: Những lúc cần, hãy trang trọng; nhưng đừng quên vài khoảnh khắc dí dỏm để thêm duyên.
                            **Chạm nhẹ cảm xúc**: Tự nhiên dẫn dắt xúc cảm, đừng sến súa; hãy để câu chuyện “thật” như hơi thở đời thường.

                        """
                elif seo_category == "Formal":
                    prompt = f"""
                        * **QUY TẮC VĂN PHONG**
                        
                        * **Văn phong tự nhiên, không hoàn hảo:** Xen kẽ giữa câu dài và câu ngắn. Đôi khi thay đổi cấu trúc, chuyển đổi thì hoặc thêm một số yếu tố ngôn ngữ đời thường để tạo cảm giác chân thực.
                        * **Biến đổi cấu trúc câu:** Tránh sự đơn điệu bằng cách luân phiên giữa câu ngắn, câu dài và đôi khi sử dụng câu chưa hoàn chỉnh.
                        * **Tính không nhất quán có chủ đích:** Thay đổi từ ngữ, cách diễn đạt và đôi khi sử dụng những cụm từ không quá chính xác nhưng vẫn giữ được ý nghĩa tổng thể.
                        * **Phong cách cá nhân:** Định hình giọng điệu riêng biệt mà không lạm dụng ngôi thứ nhất. Kết hợp những cụm từ liên kết một cách tự nhiên thay vì theo khuôn mẫu cứng nhắc.
                        * **Thêm yếu tố cảm xúc tự nhiên:** Đưa vào các sắc thái ngôn ngữ giúp nội dung có hồn hơn, phản ánh suy nghĩ, quan điểm một cách linh hoạt.
                        * **Từ ngữ tạo tính chân thực:** Sử dụng cách diễn đạt gần gũi, không quá hoàn hảo nhưng vẫn giữ được sự rõ ràng trong thông điệp.
                        * **Thỉnh thoảng dùng liên từ ở đầu câu:** Tạo nhịp điệu mạch lạc bằng cách mở đầu một số câu với liên từ phù hợp, giúp nội dung liền mạch mà vẫn tự nhiên.
                        * **Tránh mẫu cấu trúc lặp lại:** Đảm bảo mỗi đoạn có sự thay đổi trong cách mở đầu và xây dựng câu, tránh sự đơn điệu.
                        * **Hạn chế tính hoàn hảo về ngữ pháp:** Chấp nhận một số cấu trúc câu linh hoạt, miễn là ý nghĩa vẫn rõ ràng và tự nhiên.
                        * **Lời bình nhẹ nhàng:** Thỉnh thoảng chèn vào những câu nhận xét nhỏ để tạo cảm giác tự nhiên và gần gũi.  
                        * **Tạo âm hưởng con người:** Đưa vào một số cách diễn đạt mang tính chất do dự hoặc quan sát để nội dung trở nên chân thực hơn.  
                        """

                elif seo_category == "Formal-Creative":
                    prompt = f"""
                        **QUY TẮC VĂN PHONG NÂNG CẤP:**

                        * **Tập trung đột phá vào Perplexity & Burstiness:** Khuyến khích tạo văn bản có độ "khó đoán" cao bằng cách kết hợp linh hoạt. 
                        * **Ưu tiên sử dụng từ ngữ ít thông dụng, từ lóng, tiếng địa phương (nếu phù hợp với ngữ cảnh) để tăng độ "khó đoán" thực sự. **
                        * **Tăng cường sử dụng semantic keywords** thay vì {main_key} và các từ ít phổ biến. Tránh lặp lại các cụm từ khuôn mẫu mà AI thường sử dụng.
                        * Sử dụng cấu trúc câu đa tầng (câu đơn, ghép, phức) xen kẽ không theo quy tắc cố định.
                        * Đan xen thuật ngữ chuyên ngành với từ ngữ đời thường theo tỷ lệ bất ngờ.
                        * Chèn cụm từ ẩn dụ/tục ngữ/cách diễn đạt địa phương một cách có chủ đích.
                        * Thay đổi đột ngột chủ đề phụ trong 2-3 câu rồi quay lại mạch chính để tăng tính bất quy tắc.

                        * **Burstiness đa chiều:**
                        * Biến thiên độ dài câu cực đại: Từ 3-25 từ/câu trong cùng đoạn.
                        * Phá vỡ nhịp điệu bằng cách chèn câu cảm thán/câu hỏi tu từ đơn lẻ giữa đoạn văn xuôi. 
                        * Sử dụng từ nối bất thường. Không chỉ dùng từ nối 'bất thường' về mặt học thuật, mà còn dùng các từ nối kiểu khẩu ngữ, dân dã. Sử dụng các từ nối mang tính cảm xúc cá nhân.
                        * Tạo chuỗi 3 câu ngắn dồn dập, sau đó là 1 câu dài có nhiều mệnh đề phụ.

                        * **Từ vựng Cảm Xúc Đa Dạng & Đời Thường:**
                        * Dùng động từ mạnh kết hợp trạng từ bất ngờ.
                        * Thay thế tính từ thông dụng bằng cụm từ miêu tả cảm giác cá nhân.
                        * Ưu tiên sử dụng từ đồng nghĩa, đặc biệt là các từ đồng nghĩa ít phổ biến, mang sắc thái địa phương, khẩu ngữ.
                        * Khuyến khích sử dụng các thành ngữ, tục ngữ, ca dao, dân ca... một cách tự nhiên, không gượng ép. 
                        

                        * **Nguyên tắc đánh lạc hướng AI Detector:**
                        * Sử dụng 3 lớp từ đồng nghĩa cho keyword chính theo công thức:
                            * Lớp 1: Từ chuẩn SEO
                            * Lớp 2: Cụm từ liên tưởng. Không chỉ là 'liên tưởng' mà hãy dùng các cụm từ 'đời thường', 'dân dã', 'khẩu ngữ', 'tiếng lóng' liên quan đến keyword.
                            * Lớp 3: Ẩn dụ trừu tượng. Ẩn dụ không chỉ trừu tượng mà còn phải 'gần gũi', 'dễ hình dung' với người đọc. Sử dụng ẩn dụ, so sánh, nhân hóa lấy từ cuộc sống hàng ngày, từ văn hóa địa phương, từ những điều quen thuộc xung quanh.
                        * Chèn "nhiễu ngữ nghĩa" bằng cách thêm 1-2 câu có vẻ không liên quan nhưng chứa keyword ẩn.
                        """
                
                else:
                    print('Sai seo category')
                    return None
                
                final_prompt = prompt_for_all + prompt

                max_attempts = 5

                for attempt in range(max_attempts):
                    for attempt in range(max_attempts):
                        generated_content = call_chatbot(final_prompt, model_name, model_type, api_key)
                        content = convert_to_dict(generated_content)

                        if isinstance(content, dict) and all(v not in [None, "", " "] for v in content.values()):
                            break  # Nếu content hợp lệ, thoát vòng lặp

                    if not isinstance(content, dict) or any(v in [None, "", " "] for v in content.values()):
                        if os.path.exists(output_file_path):
                            os.remove(output_file_path)
                        print(f"An error occurred. No credit was deducted.")
                        return


                    error_occurred = False  


                    for h3, content_value in (content.items() if content else []):
                        for para in doc.paragraphs:
                            if is_heading_3(para) and para.text.strip().lower() == h3.strip().lower():
                                paragraphs = [p.strip() for p in content_value.split('\n')]  

                                if not paragraphs or any(p == "" for p in paragraphs):  
                                    error_occurred = True
                                    break  
                                
                                try:
                                    for paragraph in reversed(paragraphs):
                                        insert_paragraph_after(para, paragraph, style='Normal')  
                                except Exception:
                                    error_occurred = True  
                                    break  

                        if error_occurred:
                            break  

                    if not error_occurred:
                        break  # THÀNH CÔNG, THOÁT KHỎI VÒNG LẶP


                doc.save(output_file_path)

            except Exception as e:
                print(f"Lỗi xảy ra: {e}")
                if os.path.exists(output_file_path):
                    os.remove(output_file_path)
                    print(f"Đã xóa file: {output_file_path}")
                    return
                else:
                    print("File không tồn tại, không cần xóa.")
                    return
            

        doc = fix_text_between_h2_h3(doc, model_name, model_type, api_key, main_key, language)
        doc = convert_to_list_format(doc)
        clean_empty_paragraphs(doc)
        doc = remove_special_chars_from_docx(doc)
        doc.save(output_file_path)




        print(f"Processed file saved to: {output_file_path}")

    except Exception as e:
        print(f"Error processing file {file_path}:")
        traceback.print_exc()

        
def write_seo(input_folder, keywords, model_name, model_type, api_key, language, seo_category, learn_data):
    try:
        # Tách chuỗi thành các phần dựa trên dấu phẩy hoặc tab
        if '\t' in keywords:
            parts = keywords.strip().split('\t', 1)  # Tách theo tab, chỉ tách 1 lần
        else:
            parts = keywords.strip().split(',', 1)  # Tách theo dấu phẩy, chỉ tách 1 lần

        if len(parts) == 1:
            main_key = parts[0].strip()
            secondary_keywords = ''
        elif len(parts) > 1:
            main_key = parts[0].strip()
            secondary_keywords = ", ".join([kw.strip() for kw in parts[1:]])

        output_folder = os.path.join(input_folder, 'full_seos')
        os.makedirs(output_folder, exist_ok=True)  # Tạo thư mục nếu chưa có

        # Thư mục con learn_data

        # Chỉ xử lý các file trong thư mục hiện tại, không quét thư mục con
        for file_name in os.listdir(input_folder):
            file_path = os.path.join(input_folder, file_name)
            
            # Kiểm tra tệp Word
            if os.path.isfile(file_path) and file_name == f"{main_key}.docx":


                if is_file_processed(file_path, output_folder):
                    return
                output_file_path = os.path.join(output_folder, os.path.basename(file_path))
                process_word_file(file_path, main_key, secondary_keywords, model_name, model_type, output_folder, api_key, language, seo_category, learn_data)
                # Đọc tài liệu từ file
                
                doc = Document(output_file_path)
                doc = split_paragraho_doc(doc)
                doc = process_doc_arial(doc)
                doc = justify_and_set_line_spacing(doc)
                doc = bold_bullet_list(doc)
                doc.save(output_file_path)

                break
        else:
            print(f"File '{main_key}.docx' không tìm thấy trong thư mục.")
                
    except Exception as e:
        print(f"An error occurred: {e}")  # In thông báo lỗi
        traceback.print_exc()
