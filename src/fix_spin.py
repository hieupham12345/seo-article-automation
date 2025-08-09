
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import requests
import re
from docx import Document
import sys
import http.client
import json
import http.client
from docx.shared import Pt
from dotenv import load_dotenv
from src.write_seo import extract_word_to_json
from src.chat_bot import call_chatbot

# Lấy thư mục cha của thư mục chứa file code (để trỏ đúng vào .env ở folder cha)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
ENV_PATH = os.path.join(BASE_DIR, ".env")

# Load file .env từ thư mục cha
load_dotenv(ENV_PATH)






def search_query(query, api_key):
    
    query = f'"{query}"'

    conn = http.client.HTTPSConnection("google.serper.dev")
        
    payload = json.dumps({
        "q": query,
        "hl": "vi"
    })
    
    headers = {
        'X-API-KEY': api_key,  # Đảm bảo API key  
        'Content-Type': 'application/json'
    }
    
    conn.request("POST", "/search", payload, headers)
    res = conn.getresponse()
    data = res.read()
    
    # Parse the JSON response
    result = json.loads(data.decode("utf-8"))
    
    # Check if there's at least one search result in "organic" list
    if "organic" in result and len(result["organic"]) > 0:
        return False
    else:
        return True 
    



def extract_sentences_from_word(doc_path):
    # Mở file Word
    doc = Document(doc_path)
    
    # Lấy tất cả nội dung văn bản trong file
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    
    # Tách văn bản theo dấu xuống dòng (\n) thành các đoạn
    lines = full_text.split('\n')
    
    # Tách các câu từ từng dòng
    sentences = []
    for line in lines:
        # Tách câu theo các dấu câu kết thúc (dấu chấm, dấu hỏi, dấu chấm than)
        line_sentences = re.split(r'(?<=\w[.!?])\s+', line.strip())
        
        # Xử lý loại bỏ khoảng trắng thừa và kiểm tra độ dài câu
        for sentence in line_sentences:
            sentence = sentence.strip()
            if len(sentence.split()) > 6:  # Chỉ lấy câu có hơn 6 từ
                sentences.append(sentence)
    
    # Tạo tên file đầu ra (đổi phần mở rộng thành .txt)
    output_path = doc_path.replace('.docx', '.txt')
    
    # Kiểm tra nếu file tồn tại, không ghi đè
    if os.path.exists(output_path):
        print(f"File '{output_path}' đã tồn tại. Không ghi đè.")
        return
    
    # Lưu các câu vào file TXT
    with open(output_path, 'w', encoding='utf-8') as f:
        for sentence in sentences:
            f.write(sentence + '\n')
    

def process_document(doc, key_word=None):
    """
    Xử lý file docx đã mở để loại bỏ khoảng trắng ở đầu và cuối mỗi run trong đoạn văn.
    Nếu đoạn văn chứa từ khóa và đoạn văn đó là từ khóa chính xác, sẽ không xử lý đoạn văn đó.
    Chỉ xử lý các đoạn từ đoạn thứ 4 từ trên xuống đến đoạn thứ 2 từ dưới lên.

    Args:
        doc: Đối tượng Document đã được mở bằng docx.Document().
        key_word: Từ khóa để xác định những đoạn văn không cần xử lý (mặc định là None).

    Returns:
        Đối tượng Document sau khi đã xử lý, hoặc None nếu có lỗi.
    """
    try:
        # Tổng số đoạn văn trong tài liệu
        total_paragraphs = len(doc.paragraphs)

        # Xác định phạm vi các đoạn cần xử lý
        start_index = 3  # Đoạn thứ 4 từ trên xuống (chỉ số 3)
        end_index = total_paragraphs - 2  # Đoạn thứ 2 từ dưới lên

        if total_paragraphs < 4:
            print("Tài liệu không đủ đoạn văn để xử lý.")
            return doc

        # Duyệt qua các đoạn từ start_index đến end_index
        for i in range(start_index, end_index + 1):
            paragraph = doc.paragraphs[i]

            # Nếu từ khóa có và đoạn văn chính xác bằng từ khóa, bỏ qua đoạn văn này
            if key_word and paragraph.text.strip() == key_word:
                continue

            # Loại bỏ khoảng trắng ở đầu và cuối mỗi run
            for run in paragraph.runs:
                run.text = run.text.strip()  # Xóa khoảng trắng ở đầu và cuối mỗi run

            # Kiểm tra nếu đoạn văn chỉ còn khoảng trắng sau khi đã xử lý các run
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)  # Xóa đoạn văn chỉ chứa khoảng trắng

        return doc
    except Exception as e:
        print(f"Lỗi trong quá trình xử lý document: {e}")
        return None
    

def replace_sentence_in_docx(doc, original_sentence, paraphrased_sentence):
    """
    Thay thế một câu trong tài liệu DOCX, giữ nguyên format và dọn dẹp khoảng trắng dư thừa.
    Nếu đoạn là heading (h1, h2, h3) hoặc in nghiêng/căn giữa, thì bỏ dấu '.' ở cuối paraphrased_sentence.
    """
    
    for para in doc.paragraphs:
        # Kiểm tra nếu original_sentence có trong đoạn văn
        if original_sentence in para.text:
            paraphrased_sentence = re.sub(r'[\W_]+$', '', paraphrased_sentence.strip())

            # Kiểm tra xem đoạn văn có phải là heading hoặc có format đặc biệt không
            if para.style.name in ['Heading 1', 'Heading 2', 'Heading 3'] or para.alignment in [1, 2]:  # Căn giữa hoặc heading
                # Nếu là heading hoặc căn giữa, bỏ dấu '.' ở cuối
                paraphrased_sentence = paraphrased_sentence.rstrip('.')

            # Nếu original_sentence là toàn bộ đoạn văn (không phải trong một run cụ thể)
            if original_sentence == para.text:
                
                # Xác định font size dựa trên heading
                if para.style.name == 'Heading 1':
                    font_size = Pt(20)
                elif para.style.name == 'Heading 2':
                    font_size = Pt(17)
                elif para.style.name == 'Heading 3':
                    font_size = Pt(13)
                else:
                    font_size = Pt(11)

                # Lưu lại các thuộc tính format của đoạn văn
                alignment = para.alignment  # Căn lề (trái, phải, giữa, đều)
                bold = para.runs[0].bold if para.runs else False
                italic = para.runs[0].italic if para.runs else False
                underline = para.runs[0].underline if para.runs else False
                font_color = para.runs[0].font.color.rgb if para.runs else None

                # Thay thế toàn bộ đoạn văn bằng câu paraphrased, giữ nguyên các thuộc tính format
                para.clear()  # Xóa nội dung cũ của paragraph
                run = para.add_run(paraphrased_sentence)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                run.font.size = font_size
                run.font.name = "Arial"
                run.font.color.rgb = font_color

                # Đặt lại căn lề cho đoạn văn
                para.alignment = alignment
            else:
                # Lưu trữ các run mới sau khi thay thế
                new_runs = []
                for run in para.runs:
                    if original_sentence in run.text:
                        # Tách phần câu chứa original_sentence và thay thế bằng paraphrased_sentence
                        replaced_text = run.text.replace(original_sentence, paraphrased_sentence)
                        # Loại bỏ khoảng trắng dư
                        replaced_text = " ".join(replaced_text.split())

                        run.text = replaced_text
                    
                    # Lưu run vào danh sách, giữ nguyên format
                    new_runs.append(run)

                # Xóa nội dung cũ của paragraph
                para.clear()
                # Thêm lại các run đã xử lý vào paragraph
                for run in new_runs:
                    new_run = para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    new_run.font.color.rgb = run.font.color.rgb  # Giữ màu sắc của font

    return doc

def extract_sentences_from_txt(txt_file_path):
    """Trích xuất danh sách câu từ file .txt, loại bỏ các câu bắt đầu bằng '[không tìm thấy]', và trả về danh sách đảo ngược."""
    try:
        with open(txt_file_path, "r", encoding="utf-8") as file:
            # Lọc các dòng rỗng và các dòng bắt đầu bằng '[không tìm thấy]'
            sentences = [
                line.strip() for line in file 
                if line.strip() and not line.strip().startswith("[không tìm thấy]")
            ]
        return sentences[::-1]  # Đảo ngược danh sách trước khi trả về
    except FileNotFoundError:
        print(f"File .txt không tồn tại: {txt_file_path}")
        return []


def clean_text(text):
    """Loại bỏ dấu câu và khoảng trắng để so sánh chính xác hơn."""
    return re.sub(r'\s+|[^\w\s]', '', text.lower())



def normalize_text(text):
    """
    Chuẩn hóa văn bản: 
    - Dọn dẹp khoảng trắng thừa.
    - Loại bỏ khoảng trắng ở đầu và cuối.
    - Loại bỏ dấu nháy đơn ' và nháy đôi ".
    """
    # Loại bỏ các dấu nháy đơn và nháy đôi
    text = re.sub(r"[\"']", "", text)
    # Loại bỏ khoảng trắng thừa
    text = re.sub(r'\s+', ' ', text)
    # Loại bỏ khoảng trắng ở đầu và cuối chuỗi
    return text.strip()


def find_sentences(doc, search_string, exact_match=False):
    try:
        # Kiểm tra nếu doc hoặc search_string không hợp lệ
        if not doc or not hasattr(doc, 'paragraphs'):
            raise ValueError("Document không hợp lệ hoặc không có thuộc tính 'paragraphs'.")
        if not search_string or not isinstance(search_string, str):
            raise ValueError("Search string không hợp lệ, phải là chuỗi.")

        # Loại bỏ dấu và khoảng trắng ở đầu và cuối
        search_string = re.sub(r'^[.!?]+', '', search_string).strip()  
        search_string = re.sub(r'[.!?]+$', '', search_string)  # Loại bỏ dấu ở cuối

        # Chuẩn hóa chuỗi tìm kiếm để so khớp
        search_string_normalized = normalize_text(search_string.lower())

        # Danh sách lưu các câu tìm được
        result_sentences = []

        # Duyệt qua tất cả các paragraph trong document
        for para in doc.paragraphs:
            # Bỏ qua heading nếu không phải các heading cần xét
            if para.style.name.startswith('Heading'):
                text = para.text.strip()
                normalized_text = normalize_text(text.lower())  # Chuyển heading về chữ thường và chuẩn hóa
                if (exact_match and normalized_text == search_string_normalized) or (not exact_match and search_string_normalized in normalized_text):
                    result_sentences.append(text)
                    continue

            # Phân tách paragraph thành các câu theo logic giữ nguyên dấu câu
            sentences = re.split(r'(?<=[.!?])\s+', para.text)

            # Duyệt qua từng câu trong paragraph để tìm string
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence:
                    # Chuẩn hóa câu và chuyển về chữ thường để so khớp
                    normalized_sentence = normalize_text(sentence.lower())  
                    if exact_match:
                        if normalized_sentence == search_string_normalized:
                            result_sentences.append(sentence)
                    else:
                        if search_string_normalized in normalized_sentence:
                            result_sentences.append(sentence)

        # Nếu không tìm thấy câu nào, trả về thông báo
        if not result_sentences:
            return False

        # Trả về câu đầu tiên tìm được, giữ nguyên dấu câu
        return result_sentences[0].strip()

    except Exception as e:
        # Xử lý lỗi
        return f"Đã xảy ra lỗi: {str(e)}"

    

def clean_headings_in_doc(document):
    """
    Kiểm tra và làm sạch heading cấp 1, 2, và 3 trong tài liệu DOCX.
    - Heading 1: Arial, cỡ 20, in đậm, căn giữa.
    - Heading 2: Arial, cỡ 17, in đậm, màu đen.
    - Heading 3: Arial, cỡ 13, in đậm, màu đen.
    Đồng thời loại bỏ dấu '.' ở cuối heading (nếu có).
    
    Args:
        document (docx.Document): Đối tượng tài liệu đã được mở.
    
    Returns:
        docx.Document: Tài liệu đã được làm sạch và định dạng.
    """
    for paragraph in document.paragraphs:
        # Kiểm tra và loại bỏ dấu '.' ở cuối heading
        if paragraph.style.name in ["Heading 1", "Heading 2", "Heading 3"]:
            if paragraph.text.endswith('.'):
                paragraph.text = paragraph.text.rstrip('.')
            
            # Định dạng Heading 1
            if paragraph.style.name == "Heading 1":
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = None  # Mặc định là màu đen
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Căn giữa

            # Định dạng Heading 2
            elif paragraph.style.name == "Heading 2":
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(17)
                run.font.bold = True
                run.font.color.rgb = None  # Mặc định là màu đen

            # Định dạng Heading 3
            elif paragraph.style.name == "Heading 3":
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = None  # Mặc định là màu đen
    
    return document



import os
import re
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed



def process_sentence(sentence, doc, model_name, model_type, api_key, temperature, is_spineditor):
    """
    Hàm xử lý 1 câu:
      - Tìm câu đầy đủ trong file Word.
      - Nếu không tìm thấy, trả về action "not_found".
      - Nếu tìm thấy, kiểm tra với search_query (nếu spineditor=True) hoặc luôn paraphrase (nếu spineditor=False).
      - Nếu cần paraphrase: gọi call_chatbot và làm sạch kết quả, trả về action "update".
      - Nếu không cần paraphrase: trả về action "remove" (để xóa câu).
    """
    full_sentence = find_sentences(doc, sentence)
    if not full_sentence:
        print("Không tìm thấy:", sentence)
        return (sentence, "not_found", None, None)


    search_api_key = os.environ.get('ERPER_SEARCH_API_1')
    # Kiểm tra kết quả tìm kiếm (google search đã bị loại bỏ theo yêu cầu)
    result = search_query(sentence, search_api_key) if is_spineditor else False
    result = bool(result)

    if not result:
        prompt = f"""
            Paraphrase câu sau bằng **cùng 1 ngôn ngữ gốc**. Kiểm tra kỹ, **tuyệt đối cùng 1 ngôn ngữ**.
            Tránh phóng đại và có thể viết dài hơn đa 20 ký tự, thêm 2 đến 4 từ thừa. Mục đích tránh đạo văn, viết sáng tạo.
            Chỉ cung cấp kết quả tốt nhất, không có câu dẫn và câu giới thiệu, không markdown.
            Câu cần diễn giải lại: {full_sentence}.
        """
        try:
            paraphrased_sentence = call_chatbot(prompt, model_name, model_type, api_key, 1)
        except Exception as e:
            print(f"Đã xảy ra lỗi khi gọi API: {str(e)}")
            return (sentence, "error", full_sentence, None)

        # Xử lý kết quả trả về
        if paraphrased_sentence.endswith('..'):
            paraphrased_sentence = paraphrased_sentence[:-1]
        paraphrased_sentence = re.sub(r'\s{2,}', ' ', paraphrased_sentence)

        # Loại bỏ dấu chấm cuối nếu có
        full_sentence_clean = full_sentence.rstrip('.')
        print(f"Câu gốc: {full_sentence_clean}")
        print(f"Câu paraphrase: {paraphrased_sentence}")
        return (sentence, "update", full_sentence_clean, paraphrased_sentence)
    else:
        return (sentence, "remove", full_sentence, None)


def fix_spin_file(file_path, model_name, model_type, api_key, temperature, is_spineditor, search_type="er"):
    """
    Xử lý file Word và file danh sách câu theo multi-thread:
      - Chia nhỏ danh sách câu và xử lý song song từng câu.
      - Các hàm xử lý độc lập: find_sentences, search_query, call_chatbot.
      - Các thay đổi file (thay thế câu trong Word và cập nhật file TXT) được thực hiện tuần tự sau khi nhận kết quả.
    """
    while True:
        # Lấy tên file .txt từ file Word
        base_name = os.path.splitext(file_path)[0]
        txt_file_path = f"{base_name}.txt"
        if not os.path.exists(txt_file_path):
            print(f"Không tìm thấy file danh sách câu: {txt_file_path}. Bỏ qua file Word: {file_path}.")
            return

        sentences = extract_sentences_from_txt(txt_file_path)
        if not sentences:
            print(f"Danh sách câu rỗng hoặc không thể đọc từ {txt_file_path}. Dừng.")
            break

        doc = Document(file_path)
        results = []

        # Xử lý song song danh sách câu với 5 luồng
        with ThreadPoolExecutor(max_workers=5) as executor:
            future_to_sentence = {
                executor.submit(process_sentence, sentence, doc, model_name, model_type, api_key, temperature, is_spineditor): sentence
                for sentence in sentences
            }
            for future in as_completed(future_to_sentence):
                try:
                    res = future.result()
                    results.append(res)
                except Exception as exc:
                    sent = future_to_sentence[future]
                    print(f"Câu '{sent}' gây ra lỗi: {exc}")


        # Tạo map kết quả: key là câu gốc (sau khi strip), value là tuple (action, full_sentence, paraphrased_sentence)
        update_map = {}
        for original_sentence, action, full_sentence, paraphrased_sentence in results:
            update_map[original_sentence.strip()] = (action, full_sentence, paraphrased_sentence)
            # Nếu cần cập nhật trong file Word thì thực hiện thay thế ngay (theo thứ tự)
            if action == "update":
                doc = replace_sentence_in_docx(doc, full_sentence, paraphrased_sentence)

        # Lưu lại file Word sau khi đã thay thế tất cả các câu
        doc = clean_headings_in_doc(doc)
        doc.save(file_path)

        # Cập nhật file TXT dựa theo kết quả xử lý
        with open(txt_file_path, "r", encoding="utf-8") as txt_file:
            lines = txt_file.readlines()

        new_lines = []
        for line in lines:
            stripped_line = line.strip()
            if stripped_line in update_map:
                action, full_sentence, paraphrased_sentence = update_map[stripped_line]
                if action == "not_found":
                    new_lines.append(f"[không tìm thấy] {stripped_line}\n")
                elif action == "update":
                    # Nếu đang dùng spineditor thì cập nhật câu paraphrase, nếu không thì loại bỏ câu đã xử lý
                    if is_spineditor:
                        new_lines.append(normalize_text(paraphrased_sentence) + "\n")
                elif action == "remove":
                    # Loại bỏ câu khỏi file TXT
                    continue
                elif action == "error":
                    new_lines.append(line)
            else:
                new_lines.append(line)

        with open(txt_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.writelines(new_lines)

        # Nếu tất cả các dòng (không rỗng) đều đã được đánh dấu là "[không tìm thấy]", thì thoát vòng lặp.
        if all(line.strip().startswith("[không tìm thấy]") for line in new_lines if line.strip()):
            print("Tất cả các câu trong file đều không tìm thấy, thoát khỏi vòng lặp.")
            break
